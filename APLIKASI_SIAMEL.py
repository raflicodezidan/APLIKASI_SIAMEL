import os
import pandas as pd
import mysql.connector
from flask import Flask, request, render_template, redirect, url_for, flash,send_file, session,send_from_directory,make_response,jsonify
import mysql.connector
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
import openpyxl
from bs4 import BeautifulSoup
from datetime import datetime
from openpyxl.styles import Font, Alignment, PatternFill, Side, Border
from openpyxl import Workbook
from openpyxl.utils import get_column_letter
import re
from io import BytesIO
from itsdangerous import URLSafeTimedSerializer
from flask_mail import Mail, Message
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import io
from openpyxl.utils.dataframe import dataframe_to_rows
import numpy as np
import csv

app = Flask(__name__)
processed_df = None
finalized_df = None  # This will be used to store the final locked DataFrame
finalized_individu = None
finalized_institusi = None
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 587  # Gunakan 465 untuk SSL
app.config['MAIL_USE_TLS'] = True  # Atau MAIL_USE_SSL untuk port 465
app.config['MAIL_USERNAME'] = 'raflijrockdewasg@gmail.com'
app.config['MAIL_PASSWORD'] = 'jrockdewasg'
app.config['MAIL_DEFAULT_SENDER'] = 'raflijrockdewasg@gmail.com'

mail = Mail(app)

app.secret_key = 'supersecretkey'

def get_db_connection():
    return mysql.connector.connect(
        host='raflimaulana.mysql.pythonanywhere-services.com',
        user='raflimaulana',
        password='pnmim2024',
        database='raflimaulana$aml_db'  # Correct database name
    )


UPLOAD_FOLDER = '/home/raflimaulana/mysite/uploads'

# Konfigurasi folder upload
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Tentukan ekstensi file yang diizinkan
app.config['ALLOWED_EXTENSIONS'] = {'csv', 'xlsx'}
# Function to check if file format is allowed
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

# Map file keys to table names
table_mapping = {
    'individual': 'individual',
    'institusi': 'institusi',
    'dttot': 'dttot',
    'dppspm': 'dppspm',
    'judionline': 'judionline'
}

# Relevant columns for each table
# Relevant columns for each table
relevant_columns = {
    'individual': ['First Name', 'Middle Name', 'Last Name','ID No.'],
    'institusi': [
        'Authorized Person 1 First Name',
        'Authorized Person 1 Middle Name',
        'Authorized Person 1 Last Name',
        'Authorized Person 2 First Name',
        'Authorized Person 2 Middle Name',
        'Authorized Person 2 Last Name',
        'Company Name'
    ],
    'dttot': ['Nama', 'Deskripsi', 'Alamat'],  # Tambahkan kolom yang diperlukan di sini
    'dppspm': ['Nama', 'Alamat', 'Informasi Lain'],  # Tambahkan kolom yang diperlukan di sini
    'judionline': ['NAMA REKENING', 'Nomor Rekening','Bank','NIK']  # Tambahkan kolom yang diperlukan di sini
}



def detect_delimiter(file_path):
    with open(file_path, 'r') as csvfile:
        dialect = csv.Sniffer().sniff(csvfile.read(1024))
        return dialect.delimiter

# Fungsi membaca file CSV dengan delimiter yang benar
def read_csv_with_delimiter(file_path):
    delimiter = detect_delimiter(file_path)
    return pd.read_csv(file_path, delimiter=delimiter)


def save_to_database(file_path, file_key):
    table_name = table_mapping.get(file_key)

    if not table_name:
        flash(f"Tabel untuk {file_key} tidak ditemukan!", 'danger')
        return False  # Kembalikan False jika tabel tidak ditemukan

    conn = get_db_connection()
    cursor = conn.cursor()

    try:
        # Baca data dari file Excel
        if file_key == 'judionline':
            df = pd.read_excel(file_path, skiprows=[0, 1])  # Menggunakan skiprows untuk judionline
        else:
            df = pd.read_excel(file_path)
        if file_key == 'dppspm' and 'Type' in df.columns:
             df = df[~df['Type'].isin(['ORANG ATAU INDIVIDUAL', 'KORPORASI ATAU ENTITAS'])]

        df.columns = df.columns.str.strip()  # Menghilangkan spasi berlebih pada nama kolom

        # Debug: Print nama kolom yang ada pada DataFrame
        print("Kolom yang ditemukan di file:", df.columns)

        # Deteksi format kolom untuk 'judionline'
        if file_key == 'judionline':
            # Format 1: Cek kolom 'Nama Rekening', 'Nomor Rekening', 'Bank'
            if {'NAMA REKENING', 'Nomor Rekening', 'Bank'}.issubset(df.columns):
                df = df[['NAMA REKENING', 'Nomor Rekening', 'Bank']]
            # Format 2: Cek kolom 'NAMA NASABAH', 'NO. REK NASABAH', 'NIK BANK'
            elif {'NAMA NASABAH', 'NO. REK NASABAH', 'NIK BANK'}.issubset(df.columns):
                df = df[['NAMA NASABAH', 'NO. REK NASABAH', 'NIK BANK']]
                df.columns = ['NAMA REKENING', 'Nomor Rekening', 'NIK']  # Mengganti NIK BANK menjadi NIK
            else:
                flash("Format kolom di file judionline tidak dikenali.", 'danger')

        elif file_key == 'dppspm':
            # Format 1: Cek kolom 'Nama Rekening', 'Nomor Rekening', 'Bank'
            if {'Nama', 'Alamat', 'Informasi Lain'}.issubset(df.columns):
                df = df[['Nama', 'Alamat', 'Informasi Lain']]
            else:
                flash("Format kolom di file dppspm tidak dikenali.", 'danger')
                return False


        # Cek kolom yang hilang untuk format lainnya
        elif file_key in relevant_columns:
            missing_columns = [col for col in relevant_columns[file_key] if col not in df.columns]
            if missing_columns:
                flash(f"Kolom yang hilang di file {file_key}: {missing_columns}", 'danger')
                return False

            df = df[relevant_columns[file_key]]

        # Isi NaN dengan string kosong
        df = df.fillna('')


        # Tambahkan kolom timestamp
        df['upload_timestamp'] = pd.Timestamp.now()

        # Buat query SQL untuk INSERT dengan ON DUPLICATE KEY UPDATE
        cols = ', '.join([f'`{col}`' for col in df.columns])
        values = ', '.join(['%s'] * len(df.columns))
        update = ', '.join([f'`{col}` = VALUES(`{col}`)' for col in df.columns if col != 'id'])  # Sesuaikan dengan primary key
        query = f"INSERT INTO `{table_name}` ({cols}) VALUES ({values}) ON DUPLICATE KEY UPDATE {update}, upload_timestamp = NOW()"

        # Masukkan data ke database
        for _, row in df.iterrows():
            cursor.execute(query, tuple(row))

        conn.commit()
        flash(f"Data berhasil disimpan di database tabel {table_name}", 'success')
        return True

    except mysql.connector.IntegrityError as e:
        flash(f"IntegrityError saat menyimpan data: {e}", 'danger')
        return False
    except Exception as e:
        flash(f"Error saat menyimpan data ke database: {e}", 'danger')
        return False
    finally:
        cursor.close()
        conn.close()



def create_combined_key(row, keys):
    parts = []
    for key in keys:
        part = str(row[key]) if pd.notnull(row[key]) else ''
        parts.append(part)  # Selalu tambahkan bagian, termasuk yang kosong
    return ' '.join(part for part in parts if part).strip()  # Gabungkan hanya yang tidak kosong

def add_alias_columns_if_not_exist(cursor, max_aliases):
    for i in range(1, max_aliases + 1):
        column_name = f'alias{i}'
        try:
            cursor.execute(f"ALTER TABLE dttot ADD COLUMN {column_name} VARCHAR(255);")
        except Exception as e:
            # Kolom sudah ada, abaikan error
            print(f"Kolom {column_name} mungkin sudah ada, melewatkan. Error: {e}")

def insert_aliases_to_db(external_df, cursor):
    if external_df.empty:
        print("DataFrame external_df kosong, tidak ada data untuk diproses.")
        return

    if 'Nama' not in external_df.columns:
        print("Kolom 'Nama' tidak ditemukan di DataFrame.")
        return

    for index, row in external_df.iterrows():
        nama_asli = row['Nama']

        # Pastikan nama_asli tidak kosong sebelum melanjutkan
        if not nama_asli:
            print(f"Nama pada baris {index} kosong, melewatkan.")
            continue  # Lewati jika nama kosong

        # Ambil nilai alias hanya jika kolom tersebut ada
        aliases = [row[f'alias{i+1}'] for i in range(14) if f'alias{i+1}' in row]

        values = [*aliases, nama_asli]

        # Pastikan ada alias yang valid untuk di-update
        if aliases:  # Hanya eksekusi jika ada alias
            query = "UPDATE dttot SET " + ", ".join([f"alias{i+1} = %s" for i in range(len(aliases))]) + " WHERE Nama = %s"
            print(f"Mengupdate dttot untuk Nama: {nama_asli} dengan values: {values}")  # Debugging
            try:
                cursor.execute(query, values)
            except mysql.connector.errors.IntegrityError as e:
                print(f"Duplikat entry ditemukan untuk Nama: {nama_asli}, alias: {aliases}. Skipping update.")
                continue  # Jika terjadi duplikat, lewati baris ini
        else:
            print(f"Tidak ada alias untuk {nama_asli}, melewatkan update.")


def extract_aliases(nama):
    parts = [part.strip() for part in nama.split('alias')]
    return [part for part in parts if part]
def compare_data_by_columns(internal_df, external_df, internal_keys, external_key_1, cursor, is_dttot_uploaded, external_key_2=None):
    internal_df.columns = internal_df.columns.str.strip()
    external_df.columns = external_df.columns.str.strip()

    # Buat kolom kunci gabungan untuk internal_df
    internal_df['combined_key'] = internal_df.apply(create_combined_key, axis=1, keys=internal_keys).astype(str).str.strip()

    # Mengolah kolom Nama di external_df untuk mengekstrak alias
    external_df['aliases'] = external_df[external_key_1].apply(extract_aliases)

    # Mengubah kolom aliases menjadi beberapa kolom alias
    max_aliases = external_df['aliases'].apply(len).max()
    add_alias_columns_if_not_exist(cursor, max_aliases)

    for i in range(max_aliases):
        external_df[f'alias{i+1}'] = external_df['aliases'].apply(lambda x: x[i] if i < len(x) else None)

    insert_aliases_to_db(external_df, cursor)

    matched_data = pd.DataFrame()

    # Menambahkan kolom 'Penyebab' untuk mencatat alasan kecocokan
    matched_data['Penyebab'] = None

    # Melakukan perbandingan untuk setiap kolom alias
    for i in range(max_aliases):
        alias_col = f'alias{i+1}'
        if alias_col in external_df.columns:
            # Merge berdasarkan kombinasi kunci dan perbandingan antara ID No. dan NIK
            if external_key_2:
                # Merge berdasarkan kombinasi kunci dan perbandingan antara ID No. dan NIK
                temp_matched_data = pd.merge(
                    internal_df,
                    external_df,
                    how='inner',
                    left_on=['combined_key', 'ID No.'],
                    right_on=[alias_col, external_key_2]
                )
                if not temp_matched_data.empty:
                    temp_matched_data['Penyebab'] = "NAMA DAN NIK SAMA."
            else:
                temp_matched_data = pd.merge(internal_df, external_df, how='inner', left_on='combined_key', right_on=alias_col)
                if not temp_matched_data.empty:
                    temp_matched_data['Penyebab'] = 'NAMA atau ALIAS SAMA '

            matched_data = pd.concat([matched_data, temp_matched_data], ignore_index=True)

    # Juga periksa kesamaan antara NIK dan ID No. secara eksplisit
    if external_key_2:
        nik_matched_data = pd.merge(
            internal_df,
            external_df,
            how='inner',
            left_on='ID No.',
            right_on=external_key_2
        )
        if not nik_matched_data.empty:
            nik_matched_data['Penyebab'] = 'NIK SAMA '
        matched_data = pd.concat([matched_data, nik_matched_data], ignore_index=True)

    # **Menambahkan logika untuk perbandingan berdasarkan NAMA REKENING (misal di judionline)**
    if 'NAMA REKENING' in external_df.columns:
        # Mencari data yang cocok berdasarkan NAMA REKENING, tetapi menghindari duplikasi dengan hasil yang sudah ada
        rekening_matched_data = pd.merge(
            internal_df,
            external_df,
            how='inner',
            left_on='combined_key',
            right_on='NAMA REKENING'
        )
        rekening_matched_data = rekening_matched_data[~rekening_matched_data['combined_key'].isin(matched_data['combined_key'])]

    # Cek apakah matched_data tidak kosong
    if matched_data.empty:
        print("Tidak ada data yang cocok ditemukan.")  # Logging jika tidak ada data

    # Pastikan ada data yang cocok sebelum disimpan ke sesi
    if not matched_data.empty:
        session['matched_data'] = matched_data.to_html(index=False, classes='table table-bordered table-striped')  # Simpan data ke sesi
        print("Data matched disimpan ke sesi.")  # Logging

    # Ganti nama kolom combined_key menjadi Kunci Penggabungan
    if 'combined_key' in matched_data.columns:
        matched_data.rename(columns={'combined_key': 'Nama Nasabah PNM'}, inplace=True)

    # Hapus kolom aliases dan alias dari hasil akhir jika ada
    if 'aliases' in matched_data.columns:
        matched_data.drop(columns=['aliases'], inplace=True)
    for i in range(max_aliases):
        alias_col = f'alias{i+1}'
        if alias_col in matched_data.columns:
            matched_data.drop(columns=[alias_col], inplace=True)

    return matched_data.to_html(index=False, classes='table table-bordered table-striped') if not matched_data.empty else "<p>Tidak ada data yang cocok.</p>"


# Fungsi utama untuk perbandingan data
def compare_data(uploaded_files):
    conn = get_db_connection()
    cursor = conn.cursor()

    try:
        tables = {
            'individual': 'individual',
            'institusi': 'institusi',
            'dttot': 'dttot',
            'dppspm': 'dppspm',
            'judionline': 'judionline'
        }

        data_frames = {}
        latest_timestamp_query = "SELECT MAX(upload_timestamp) FROM {} WHERE upload_timestamp IS NOT NULL"

        for table, key in tables.items():
            if key in uploaded_files:  # Cek apakah file ini diunggah
                # Dapatkan timestamp terbaru dari tabel
                cursor.execute(latest_timestamp_query.format(table))
                latest_timestamp = cursor.fetchone()[0]

                # Jika tidak ada timestamp (belum ada data diunggah), skip tabel ini
                if latest_timestamp is None:
                    continue

                # Ambil data dengan timestamp terbaru dan lakukan pengecekan untuk data baru
                cursor.execute(f"SELECT * FROM {table} WHERE upload_timestamp = %s", (latest_timestamp,))
                df = pd.DataFrame(cursor.fetchall(), columns=[i[0] for i in cursor.description])

                if df.empty:
                    continue

                df = df[relevant_columns[key]]  # Filter ke kolom yang relevan
                data_frames[table] = df

        matched_data = {}

        # Lakukan perbandingan data hanya jika file terbaru yang relevan ada
        if 'individual' in uploaded_files:
            if 'dppspm' in uploaded_files:
                matched_data['DATA INDIVIDUAL - DPPSPM'] = compare_data_by_columns(
                    data_frames['individual'],
                    data_frames['dppspm'],
                    ['First Name', 'Middle Name', 'Last Name'],
                    'Nama',
                    cursor,
                    False
                )

            if 'dttot' in uploaded_files:
                matched_data['DATA INDIVIDUAL - DTTOT'] = compare_data_by_columns(
                    data_frames['individual'],
                    data_frames['dttot'],
                    ['First Name', 'Middle Name', 'Last Name'],
                    'Nama',
                    cursor,
                    True
                )

            if 'judionline' in uploaded_files:
                matched_data['DATA INDIVIDUAL - JUDI ONLINE'] = compare_data_by_columns(
                    data_frames['individual'],
                    data_frames['judionline'],
                    ['First Name', 'Middle Name', 'Last Name'],
                    'NAMA REKENING',
                    cursor,
                    False,
                    'NIK'  # Tambahkan parameter NIK untuk perbandingan dengan ID No.
                )

        if 'institusi' in uploaded_files:
            if 'dppspm' in uploaded_files:
                matched_data['DATA INSTITUSI - DPPSPM'] = compare_data_by_columns(
                    data_frames['institusi'],
                    data_frames['dppspm'],
                    ['Authorized Person 1 First Name', 'Authorized Person 1 Middle Name', 'Authorized Person 1 Last Name',
                     'Authorized Person 2 First Name', 'Authorized Person 2 Middle Name', 'Authorized Person 2 Last Name',
                     'Company Name'],
                    'Nama',
                    cursor,
                    False
                )

            if 'dttot' in uploaded_files:
                matched_data['DATA INSTITUSI - DTTOT'] = compare_data_by_columns(
                    data_frames['institusi'],
                    data_frames['dttot'],
                    ['Authorized Person 1 First Name', 'Authorized Person 1 Middle Name', 'Authorized Person 1 Last Name',
                     'Authorized Person 2 First Name', 'Authorized Person 2 Middle Name', 'Authorized Person 2 Last Name',
                     'Company Name'],
                    'Nama',
                    cursor,
                    True
                )

            if 'judionline' in uploaded_files:
                matched_data['DATA INSTITUSI - JUDI ONLINE '] = compare_data_by_columns(
                    data_frames['institusi'],
                    data_frames['judionline'],
                    ['Authorized Person 1 First Name', 'Authorized Person 1 Middle Name', 'Authorized Person 1 Last Name',
                     'Authorized Person 2 First Name', 'Authorized Person 2 Middle Name', 'Authorized Person 2 Last Name',
                     'Company Name'],
                    'NAMA REKENING',
                    cursor,
                    False
                )

        return matched_data
    finally:
        cursor.close()
        conn.close()




@app.route('/upload', methods=['GET', 'POST'])
def upload_files():
    # Cek apakah user sudah login
    if 'username' not in session and 'admin_username' not in session:
        flash("Silakan login terlebih dahulu untuk mengakses halaman ini", "warning")
        return redirect(url_for('landingpage'))

    if request.method == 'POST':
        session.pop('matched_data', None)
        username = session.get('username') or session.get('admin_username', 'Unknown User')
        role = 'user' if 'username' in session else 'admin'
        log_activity(username, role, "Compare Data Internal - SIGAP (DTTOT, DPPSPM,FATF, & JUDI ONLINE)")

        files = {
            'dttot': request.files.get('dttot'),
            'dppspm': request.files.get('dppspm'),
            'judionline': request.files.get('judionline'),
            'individual': request.files.get('individual'),
            'institusi': request.files.get('institusi'),
        }

        uploaded_files = []

        for file_key, file in files.items():
            if file:
                if not allowed_file(file.filename):
                    flash(f"Format file {file.filename} harus .xlsx atau .csv!", 'danger')
                    return redirect(request.url)

                filename = secure_filename(file.filename)
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)

                if not os.path.exists(app.config['UPLOAD_FOLDER']):
                    os.makedirs(app.config['UPLOAD_FOLDER'])

                try:
                    file.save(file_path)
                    # Panggil fungsi save_to_database dan periksa hasilnya
                    if not save_to_database(file_path, file_key):
                        flash(f"Proses penyimpanan data dari {file_key} dibatalkan karena kolom yang hilang.", 'danger')
                        return redirect(request.url)
                    uploaded_files.append(file_key)  # Menyimpan file yang diunggah
                except Exception as e:
                    flash(f"Error saat memproses file {file_key}: {e}", 'danger')
                    return redirect(request.url)


        # Lakukan perbandingan setelah semua file diunggah
        matched_data = compare_data(uploaded_files)

        # Simpan data hasil perbandingan dalam sesi
        session['matched_data'] = matched_data

        return redirect(url_for('compare_results'))

    return render_template('upload.html')


@app.route('/results')
def compare_results():
    if 'username' not in session and 'admin_username' not in session:
        flash("Silakan login terlebih dahulu untuk mengakses halaman ini", "warning")
        return redirect(url_for('landingpage'))
    # Ambil data hasil perbandingan dari sesi
    matched_data = session.get('matched_data', None)
    if matched_data is None:
        flash('Tidak ada data untuk ditampilkan.', 'warning')
        return redirect(url_for('upload_files'))

    return render_template('results.html', data=matched_data)


from flask import send_file, flash, redirect, url_for, session
from openpyxl import load_workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter
import pandas as pd
import os

def adjust_column_width(sheet):
    for col in sheet.columns:
        max_length = 0
        col_name = get_column_letter(col[0].column)
        for cell in col:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(cell.value)
            except:
                pass
        adjusted_width = (max_length + 2) * 1.1
        sheet.column_dimensions[col_name].width = adjusted_width

def get_columns_by_title(title):
    if "INDIVIDUAL" in title:
        return ['First Name', 'Middle Name', 'Last Name', 'ID No.', 'Penyebab']
    elif "INSTITUSI" in title:
        return ['Company Name', 'Authorized Person 1 First Name', 'Authorized Person 1 Middle Name',
                   'Authorized Person 1 Last Name', 'Authorized Person 2 First Name',
                   'Authorized Person 2 Middle Name', 'Authorized Person 2 Last Name', 'Penyebab']
    else:
        return []

@app.route('/download/excel')
def download_excel():
    if 'username' not in session and 'admin_username' not in session:
        flash("Silakan login terlebih dahulu untuk mengakses halaman ini", "warning")
        return redirect(url_for('landingpage'))

    username = session.get('username') or session.get('admin_username', 'Unknown User')
    role = 'user' if 'username' in session else 'admin'
    log_activity(username, role, "Mengunduh Excel SIGAP (TPPU, TPPT, PPSPM, & JUDI ONLINE)")

    matched_data = session.get('matched_data')
    downloader_name = username

    if matched_data is None or not matched_data:
        flash('Tidak ada data untuk diunduh.', 'warning')
        return redirect(url_for('upload_files'))

    file_path = 'Laporan Hasil Perbandingan data Internal dengan SIGAP .xlsx'

    with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
        for title, data in matched_data.items():
            try:
                df = pd.read_html(data, flavor='bs4')[0]
            except (ValueError, IndexError):
                columns = get_columns_by_title(title)
                df = pd.DataFrame(columns=columns)
                df.loc[0] = ['NIHIL'] * len(columns)

            df.to_excel(writer, sheet_name=title[:31], index=False, startrow=1)

    wb = load_workbook(file_path)
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        dynamic_title = f"Laporan Hasil Perbandingan Data Internal & SIGAP ( {sheet_name})"

        # Title styling and merge
        sheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=sheet.max_column)
        title_cell = sheet.cell(row=1, column=1)
        title_cell.value = dynamic_title
        title_cell.alignment = Alignment(horizontal='center')
        title_cell.font = Font(size=18, bold=True)
        title_cell.fill = PatternFill(start_color='ADD8E6', end_color='ADD8E6', fill_type='solid')

        adjust_column_width(sheet)

        # Styling data cells: border, background color, and center alignment
        thin_border = Border(left=Side(style='thin', color='000000'), right=Side(style='thin', color='000000'),
                             top=Side(style='thin', color='000000'), bottom=Side(style='thin', color='000000'))
        gray_fill = PatternFill(start_color="D3D3D3", end_color="D3D3D3", fill_type="solid")

        for row in sheet.iter_rows(min_row=2, max_row=sheet.max_row, min_col=1, max_col=sheet.max_column):
            for cell in row:
                cell.border = thin_border
                cell.fill = gray_fill
                cell.alignment = Alignment(horizontal='center', vertical='center')
                if cell.value == 'NIHIL':
                    cell.fill = PatternFill(start_color='FF0000', end_color='FF0000', fill_type='solid')
                    cell.font = Font(color="FFFFFF")

        # Footer for downloader info with blue background and centered
        footer_row = sheet.max_row + 2
        sheet.merge_cells(start_row=footer_row, start_column=1, end_row=footer_row, end_column=sheet.max_column)
        footer_cell = sheet.cell(row=footer_row, column=1)
        footer_cell.value = f"Diunduh oleh: {downloader_name} - {role.capitalize()}"
        footer_cell.alignment = Alignment(horizontal='center', vertical='center')
        footer_cell.font = Font(bold=True, color="FFFFFF")
        footer_cell.fill = PatternFill(start_color="0000FF", end_color="0000FF", fill_type="solid")

    wb.save(file_path)

    return send_file(file_path, as_attachment=True)

from docx import Document
from docx.shared import RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from datetime import datetime
from io import BytesIO
from bs4 import BeautifulSoup

def create_word(matched_data, downloader_name):
    # Membuat dokumen Word baru
    doc = Document()

    # Waktu dan informasi downloader
    current_date = datetime.now().strftime('%d-%m-%Y %H:%M')

    # Menambahkan header di awal halaman pertama setiap perbandingan
    def add_header():
        # Menambahkan judul "Laporan SIGAP"
        header_paragraph = doc.add_paragraph("Laporan SIGAP")
        header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = header_paragraph.runs[0]
        run.bold = True
        run.font.size = 210000  # Ukuran font (18 pt)
        run.font.color.rgb = RGBColor(0, 0, 255)  # Warna biru

        # Menambahkan subjudul "(DTTOT, DPPSPM, FATF, & JUDI ONLINE)"
        subheader_paragraph = doc.add_paragraph("(DTTOT, DPPSPM, FATF, & JUDI ONLINE)")
        subheader_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = subheader_paragraph.runs[0]
        run.bold = True
        run.font.size = 190000  # Ukuran font (16 pt)
        run.font.color.rgb = RGBColor(0, 0, 255)  # Warna biru

        # Menambahkan tanggal dan downloader di bawah subjudul
        date_paragraph = doc.add_paragraph(f'Tanggal: {current_date} - Diunduh oleh: {downloader_name}')
        date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = date_paragraph.runs[0]
        run.font.color.rgb = RGBColor(0, 0, 255)  # Warna biru

        doc.add_paragraph('-' * 100).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for key, html_table in matched_data.items():
        if doc.paragraphs:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        add_header()  # Menambahkan header dengan tanggal dan downloader

        heading_paragraph = doc.add_paragraph(key)
        heading_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = heading_paragraph.runs[0]
        run.bold = True
        run.font.size = 190000  # Ukuran font (18 pt)
        run.font.color.rgb = RGBColor(0, 0, 255)  # Warna biru

        # Parsing tabel HTML
        soup = BeautifulSoup(html_table, 'html.parser')
        rows = soup.find_all('tr')

        # Menentukan kolom sesuai dengan key, tanpa "Penyebab" jika tidak ada data
        if len(rows) == 0:
            if key == 'DATA INDIVIDUAL - DPPSPM':
                columns = ['Nama Nasabah', 'Nama']
            elif key == 'DATA INDIVIDUAL - DTTOT':
                columns = ['ID No.', 'Nama Nasabah', 'Nama']
            elif key == 'DATA INDIVIDUAL - JUDI ONLINE':
                columns = ['Nama Nasabah', 'Nama Rekening', 'Nomor Rekening', 'NIK']
            elif key == 'DATA INSTITUSI - DPPSPM':
                columns = ['Company Name', 'Nama Nasabah', 'Nama']
            elif key == 'DATA INSTITUSI - DTTOT':
                columns = ['Company Name', 'Nama Nasabah', 'Nama']
            elif key == 'DATA INSTITUSI - JUDI ONLINE':
                columns = ['Company Name', 'Nama Nasabah', 'Nama Rekening']
        else:
            if key == 'DATA INDIVIDUAL - DPPSPM':
                columns = ['Penyebab', 'Nama Nasabah', 'Nama']
            elif key == 'DATA INDIVIDUAL - DTTOT':
                columns = ['Penyebab', 'ID No.', 'Nama Nasabah', 'Nama']
            elif key == 'DATA INDIVIDUAL - JUDI ONLINE':
                columns = ['Penyebab', 'Nama Nasabah', 'Nama Rekening', 'Nomor Rekening', 'NIK']
            elif key == 'DATA INSTITUSI - DPPSPM':
                columns = ['Penyebab', 'Company Name', 'Nama Nasabah', 'Nama']
            elif key == 'DATA INSTITUSI - DTTOT':
                columns = ['Penyebab', 'Company Name', 'Nama Nasabah', 'Nama']
            elif key == 'DATA INSTITUSI - JUDI ONLINE':
                columns = ['Penyebab', 'Company Name', 'Nama Nasabah', 'Nama Rekening']

        table = doc.add_table(rows=1, cols=len(columns), style='Table Grid')
        header_cells = table.rows[0].cells

        for i, col_name in enumerate(columns):
            header_cells[i].text = col_name

        if len(rows) == 0:
            row_cells = table.add_row().cells
            for cell in row_cells:
                cell.text = "NIHIL"
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255, 0, 0)  # Mengubah warna menjadi merah
        else:
            for row in rows:
                row_cells = table.add_row().cells
                cells = row.find_all('td')
                for i, cell in enumerate(cells):
                    if i < len(row_cells):
                        row_cells[i].text = cell.get_text().strip()
                    else:
                        break

        for i, column in enumerate(table.columns):
            for cell in column.cells:
                cell.width = Inches(4)


        doc.add_paragraph()  # Paragraf kosong untuk memberi jarak
        footer_text = f'SI AMEL PNMIM (Sistem Informasi Anti Money Laundering PNMIM)'
        footer_paragraph = doc.add_paragraph(footer_text)
        footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph()  # Menambahkan jarak sebelum halaman berikutnya

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    return output


@app.route('/download_word')
def download_word():
    if 'username' not in session and 'admin_username' not in session:
        flash("Silakan login terlebih dahulu untuk mengakses halaman ini", "warning")
        return redirect(url_for('landingpage'))

    username = session.get('username') or session.get('admin_username', 'Unknown User')
    role = 'user' if 'username' in session else 'admin'
    log_activity(username, role, " Download WORD Laporan SIGAP (DTTOT, DPPSPM,FATF, & JUDI ONLINE)")

    matched_data = session.get('matched_data', None)  # Panggil fungsi perbandingan data
    downloader_name = session.get('username') or session.get('admin_username', 'Unknown User')  # Ambil nama downloader

    word_output = create_word(matched_data, downloader_name)  # Tambahkan downloader_name

    return send_file(word_output, download_name="Laporan SIGAP (DTTOT, DPPSPM,FATF, & JUDI ONLINE).docx", as_attachment=True, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

def word_nihil(matched_data, downloader_name):
    # Membuat dokumen Word baru
    doc = Document()

    # Waktu dan informasi downloader
    current_date = datetime.now().strftime('%d-%m-%Y %H:%M')

    # Menambahkan header di awal halaman pertama setiap perbandingan
    def add_header():
        # Menambahkan judul "Laporan SIGAP"
        header_paragraph = doc.add_paragraph("Laporan SIGAP")
        header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = header_paragraph.runs[0]
        run.bold = True
        run.font.size = 210000  # Ukuran font (18 pt)
        run.font.color.rgb = RGBColor(0, 0, 255)  # Warna biru

        # Menambahkan subjudul "(DTTOT, DPPSPM, FATF, & JUDI ONLINE)"
        subheader_paragraph = doc.add_paragraph("(DTTOT, DPPSPM, FATF, & JUDI ONLINE)")
        subheader_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = subheader_paragraph.runs[0]
        run.bold = True
        run.font.size = 190000  # Ukuran font (16 pt)
        run.font.color.rgb = RGBColor(0, 0, 255)  # Warna biru

        # Menambahkan tanggal dan downloader di bawah subjudul
        date_paragraph = doc.add_paragraph(f'Tanggal: {current_date} - Diunduh oleh: {downloader_name}')
        date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = date_paragraph.runs[0]
        run.font.color.rgb = RGBColor(0, 0, 255)  # Warna biru

        doc.add_paragraph('-' * 100).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for key, html_table in matched_data.items():
        if doc.paragraphs:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        add_header()  # Menambahkan header dengan tanggal dan downloader

        heading_paragraph = doc.add_paragraph(key)
        heading_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = heading_paragraph.runs[0]
        run.bold = True
        run.font.size = 190000  # Ukuran font (18 pt)
        run.font.color.rgb = RGBColor(0, 0, 255)  # Warna biru

        # Menentukan kolom sesuai dengan key
        if key == 'DATA INDIVIDUAL - DPPSPM':
            columns = ['Penyebab', 'Nama Nasabah', 'Nama']
        elif key == 'DATA INDIVIDUAL - DTTOT':
            columns = ['Penyebab', 'ID No.', 'Nama Nasabah', 'Nama']
        elif key == 'DATA INDIVIDUAL - JUDI ONLINE':
            columns = ['Penyebab', 'Nama Nasabah', 'Nama Rekening', 'Nomor Rekening', 'NIK']
        elif key == 'DATA INSTITUSI - DPPSPM':
            columns = ['Penyebab', 'Company Name', 'Nama Nasabah', 'Nama']
        elif key == 'DATA INSTITUSI - DTTOT':
            columns = ['Penyebab', 'Company Name', 'Nama Nasabah', 'Nama']
        elif key == 'DATA INSTITUSI - JUDI ONLINE':
            columns = ['Penyebab', 'Company Name', 'Nama Nasabah', 'Nama Rekening']

        table = doc.add_table(rows=1, cols=len(columns), style='Table Grid')
        header_cells = table.rows[0].cells

        for i, col_name in enumerate(columns):
            header_cells[i].text = col_name

        # Menambahkan baris "NIHIL"
        row_cells = table.add_row().cells
        for cell in row_cells:
            cell.text = "NIHIL"
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 0, 0)  # Mengubah warna menjadi merah

        for i, column in enumerate(table.columns):
            for cell in column.cells:
                cell.width = Inches(4)

        doc.add_paragraph()  # Paragraf kosong untuk memberi jarak
        footer_text = f'SI AMEL PNMIM (Sistem Informasi Anti Money Laundering PNMIM)'
        footer_paragraph = doc.add_paragraph(footer_text)
        footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph()  # Menambahkan jarak sebelum halaman berikutnya

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    return output


@app.route('/download_word_nihil')
def download_word_nihil():
    if 'username' not in session and 'admin_username' not in session:
        flash("Silakan login terlebih dahulu untuk mengakses halaman ini", "warning")
        return redirect(url_for('landingpage'))

    username = session.get('username') or session.get('admin_username', 'Unknown User')
    role = 'user' if 'username' in session else 'admin'
    log_activity(username, role, " Download WORD Laporan SIGAP (DTTOT, DPPSPM,FATF, & JUDI ONLINE)")

    matched_data = session.get('matched_data', None)  # Panggil fungsi perbandingan data
    downloader_name = session.get('username') or session.get('admin_username', 'Unknown User')  # Ambil nama downloader

    word_output = word_nihil(matched_data, downloader_name)  # Tambahkan downloader_name

    return send_file(word_output, download_name="Laporan SIGAP (DTTOT, DPPSPM,FATF, & JUDI ONLINE).docx", as_attachment=True, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/')
def landingpage():
    return render_template('landing_page.html')


@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form['username']
        email = request.form['email']
        nomorhp = request.form['nomorhp']
        password = request.form['password']
        confirm_password = request.form['confirm_password']

        # Validasi password dan konfirmasi password
        if password != confirm_password:
            flash('Password dan konfirmasi password tidak cocok.', 'error')
            return redirect(url_for('register'))

        # Validasi email
        if not validate_email(email):
            flash('Format email tidak valid.', 'error')
            return redirect(url_for('register'))

        # Enkripsi password
        hashed_password = generate_password_hash(password)

        # Gunakan prepared statement untuk mencegah SQL injection
        conn = get_db_connection()
        cursor = conn.cursor()
        query = "INSERT INTO users (username, email, password, nomorhp) VALUES (%s, %s, %s, %s)"
        try:
            cursor.execute(query, (username, email, hashed_password, nomorhp))
            conn.commit()

            # Flash pesan keberhasilan, tapi tidak melakukan redirect langsung ke login
            flash('Registrasi berhasil! Silakan login untuk melanjutkan.', 'success')
            return render_template('register.html')  # Tetap render register page agar modal muncul
        except mysql.connector.Error as err:
            flash(f'Error: {err}', 'error')
        finally:
            cursor.close()
            conn.close()

    return render_template('register.html')


def validate_email(email):
    import re
    email_regex = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    return re.match(email_regex, email) is not None

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        remember = request.form.get('remember')  # Ambil nilai checkbox "Ingatkan Saya"

        # Koneksi ke database
        conn = get_db_connection()
        cursor = conn.cursor(dictionary=True)

        try:
            # Hapus catatan login yang lebih dari 30 hari
            cursor.execute("DELETE FROM login_records WHERE login_time < NOW() - INTERVAL 30 DAY")
            conn.commit()

            # Cek username di database
            query = "SELECT * FROM users WHERE username = %s"
            cursor.execute(query, (username,))
            user = cursor.fetchone()

            if user and check_password_hash(user['password'], password):
                # Set session dan simpan catatan login
                session['username'] = user['username']
                cursor.execute("INSERT INTO login_records (username, role) VALUES (%s, 'user')", (username,))
                conn.commit()

                # Buat respons dengan cookie jika "Ingatkan Saya" dicentang
                resp = make_response(redirect(url_for('dashboard')))
                if remember:  # Simpan username di cookie untuk 30 hari
                    resp.set_cookie('username', username, max_age=30*24*60*60)  # 30 hari
                else:  # Hapus cookie jika tidak dicentang
                    resp.set_cookie('username', '', max_age=0)

                return resp
            else:
                flash('Username atau password salah!', 'error')

        except Exception as e:
            flash(f'Terjadi kesalahan: {e}', 'error')

        finally:
            cursor.close()
            conn.close()

    # Ambil username dari cookie jika ada
    saved_username = request.cookies.get('username')
    return render_template('login.html', saved_username=saved_username)




@app.route('/dashboard')
def dashboard():
    if 'username' not in session:
        flash('Silakan login terlebih dahulu.', 'warning')
        return redirect(url_for('landingpage'))  # Arahkan ke halaman login jika tidak ada sesi

    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    # Mengambil data pengguna berdasarkan username dari sesi
    query = "SELECT username, email, nomorhp FROM users WHERE username = %s"
    cursor.execute(query, (session['username'],))
    user = cursor.fetchone()

    cursor.close()
    conn.close()

    if user:
        return render_template('dashboard.html', user=user)  # Kirim data pengguna ke template
    else:
        flash('Data pengguna tidak ditemukan.', 'danger')
        return redirect(url_for('landingpage'))  # Arahkan ke halaman login jika pengguna tidak ditemukan

@app.route('/admin/login', methods=['GET', 'POST'])
def admin_login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        remember = request.form.get('remember')  # Ambil nilai checkbox "Ingatkan Saya"

        # Koneksi ke database
        conn = get_db_connection()
        cursor = conn.cursor(dictionary=True)

        # Hapus catatan login yang lebih dari 30 hari
        cursor.execute("DELETE FROM login_records WHERE login_time < NOW() - INTERVAL 30 DAY")
        conn.commit()

        # Cek username di database
        query = "SELECT * FROM admins WHERE username = %s"
        cursor.execute(query, (username,))
        admin = cursor.fetchone()

        if admin and check_password_hash(admin['password'], password):
            # Set session dan simpan catatan login
            session['admin_username'] = admin['username']
            cursor.execute("INSERT INTO login_records (username, role) VALUES (%s, 'admin')", (username,))
            conn.commit()

            # Buat respons dengan cookie jika "Ingatkan Saya" dicentang
            resp = make_response(redirect(url_for('admin_dashboard')))
            if remember:  # Simpan username di cookie untuk 30 hari
                resp.set_cookie('username', username, max_age=30*24*60*60)  # 30 hari
            else:  # Hapus cookie jika tidak dicentang
                resp.set_cookie('username', '', max_age=0)
            cursor.close()
            conn.close()
            return resp
        else:
            flash('Username atau password salah!', 'error')

        cursor.close()
        conn.close()

    # Ambil username dari cookie jika ada
    saved_username = request.cookies.get('username')
    return render_template('admin_login.html', saved_username=saved_username)




def get_login_records():
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    # Mengambil semua catatan login dan logout, urutkan berdasarkan waktu login terbaru
    query = """
    SELECT username, role, login_time, logout_time
    FROM login_records
    ORDER BY login_time DESC
    """
    cursor.execute(query)
    records = cursor.fetchall()

    cursor.close()
    conn.close()

    return records



def get_activity_logs():
    db_connection = get_db_connection()
    cursor = db_connection.cursor(dictionary=True)

    query = "SELECT * FROM activity_logs ORDER BY timestamp DESC"
    cursor.execute(query)
    activity_logs = cursor.fetchall()

    cursor.close()
    db_connection.close()
    return activity_logs

def log_activity(username, role, activity):
    db_connection = get_db_connection()
    cursor = db_connection.cursor()

    query = """
    INSERT INTO activity_logs (username, role, activity)
    VALUES (%s, %s, %s)
    """
    cursor.execute(query, (username, role, activity))
    db_connection.commit()

    cursor.close()
    db_connection.close()

def cleanup_old_activity_logs():
    db_connection = get_db_connection()
    cursor = db_connection.cursor()

    query = """
    DELETE FROM activity_logs
    WHERE timestamp < NOW() - INTERVAL 30 DAY
    """
    cursor.execute(query)
    db_connection.commit()

    cursor.close()
    db_connection.close()

def cleanup_old_login_records():
    db_connection = get_db_connection()
    cursor = db_connection.cursor()

    query = """
    DELETE FROM login_records
    WHERE login_time < NOW() - INTERVAL 30 DAY
       OR logout_time < NOW() - INTERVAL 30 DAY
    """
    cursor.execute(query)
    db_connection.commit()

    cursor.close()
    db_connection.close()


@app.route('/admin/dashboard')
def admin_dashboard():
    if 'admin_username' not in session:
        flash('Silakan login sebagai admin terlebih dahulu.', 'warning')
        return redirect(url_for('admin_login'))  # Arahkan ke halaman login admin
    return render_template('admin_dashboard.html', admin_username=session['admin_username'])


@app.route('/logoutAdmin')
def logoutadmin():
    if 'admin_username' in session:
        # Mendapatkan username admin sebelum logout
        admin_username = session['admin_username']

        # Update waktu logout di login_records
        conn = get_db_connection()
        cursor = conn.cursor()
        query = "UPDATE login_records SET logout_time = NOW() WHERE username = %s AND logout_time IS NULL"
        cursor.execute(query, (admin_username,))
        conn.commit()
        cursor.close()
        conn.close()

        # Catat aktivitas logout

    session.pop('admin_username', None)
    session.clear()
    return redirect(url_for('landingpage'))


@app.route('/logout')
def logout():
    if 'username' in session:
        # Mendapatkan username sebelum logout
        username = session['username']

        # Update waktu logout di login_records
        conn = get_db_connection()
        cursor = conn.cursor()
        query = "UPDATE login_records SET logout_time = NOW() WHERE username = %s AND logout_time IS NULL"
        cursor.execute(query, (username,))
        conn.commit()
        cursor.close()
        conn.close()



    session.pop('username', None)
    session.clear()
    return redirect(url_for('landingpage'))





def add_admin(username, password):
    # Validasi input username dan password
    if not username or not password:
        print("Username atau password tidak boleh kosong.")
        return

    print(f'Menambahkan admin dengan username: {username}')

    hashed_password = generate_password_hash(password)
    print(f'Password yang dienkripsi: {hashed_password}')

    conn = get_db_connection()
    cursor = conn.cursor()

    # Cek apakah username sudah ada
    cursor.execute("SELECT * FROM admins WHERE username = %s", (username,))
    existing_admin = cursor.fetchone()
    if existing_admin:
        print(f'Admin dengan username {username} sudah ada.')
        cursor.close()
        cursor.execute("SELECT * FROM admins WHERE username = %s AND password = %s", (username, password))

        conn.close()
        return

    # Insert data admin baru
    query = "INSERT INTO admins (username, password) VALUES (%s, %s)"
    try:
        cursor.execute(query, (username, hashed_password))
        conn.commit()
        print(f'Admin {username} berhasil ditambahkan.')
    except mysql.connector.Error as err:
        if err.errno == mysql.connector.errorcode.ER_DUP_ENTRY:
            print(f"Error: Admin dengan username {username} sudah ada.")
        else:
            print(f'Error: {err}')
    finally:
        cursor.close()
        conn.close()






@app.route('/manage_user', methods=['GET', 'POST'])
def manage_user():
    if 'admin_username' not in session:
        flash('Silakan login sebagai admin terlebih dahulu.', 'warning')
        return redirect(url_for('admin_login'))  # Arahkan ke halaman login admin jika tidak ada sesi

    if request.method == 'POST':
        # Mendapatkan admin yang sedang login untuk mencatat aktivitas
        current_admin = session.get('admin_username')

        # Menangani form untuk menambah admin atau user
        if 'add_admin' in request.form:
            username = request.form['admin_username']
            password = request.form['admin_password']
            add_admin(username, password)
            flash(f'Admin {username} berhasil ditambahkan!', 'success')

            # Tambahkan ke activity log dengan detail admin yang ditambah
            log_activity(current_admin, 'admin', f'Add Admin : {username}')

        elif 'add_user' in request.form:
            username = request.form['user_username']
            email = request.form['user_email']
            nomorhp = request.form['user_nomorhp']
            password = request.form['user_password']
            add_user(username, email, password, nomorhp)
            flash(f'User {username} berhasil ditambahkan!', 'success')

            # Tambahkan ke activity log dengan detail user yang ditambah
            log_activity(current_admin, 'admin', f'Add User : {username}')

        elif 'delete_admin' in request.form:
            username = request.form['delete_admin_username']
            delete_admin(username)
            flash(f'Admin {username} berhasil dihapus!', 'success')

            # Tambahkan ke activity log dengan detail admin yang dihapus
            log_activity(current_admin, 'admin', f'Delete Admin : {username}')

        elif 'delete_user' in request.form:
            username = request.form['delete_user_username']
            delete_user(username)
            flash(f'User {username} berhasil dihapus!', 'success')

            # Tambahkan ke activity log dengan detail user yang dihapus
            log_activity(current_admin, 'admin', f'Delete User : {username}')

    # Ambil daftar user dan admin dari database
    admins = get_all_admins()
    users = get_all_users()

    # Ambil catatan login dan aktivitas
    login_records = get_login_records()
    activity_logs = get_activity_logs()  # Ambil catatan aktivitas dari database

    return render_template('manage_user.html', admins=admins, users=users, login_records=login_records, activity_logs=activity_logs)

# Route untuk Panduan Pengguna
@app.route('/user_guide')
def user_guide():
    if 'username' not in session and 'admin_username' not in session:
        flash("Silakan login terlebih dahulu untuk mengakses halaman ini", "warning")
        return redirect(url_for('landingpage'))
    username = session.get('username') or session.get('admin_username', 'Unknown User')
    role = 'user' if 'username' in session else 'admin'
    log_activity(username, role, "Membuka Buku Panduan Pengguna ")
    return render_template('user_guide.html')

@app.route('/download-guide')
def download_guide():
    return send_from_directory(directory='static', path='Buku Panduan Penggunaan APK Anti Money Laundering.docx', as_attachment=True)


# Route untuk FAQ
@app.route('/faq')
def faq():
    if 'username' not in session and 'admin_username' not in session:
        flash("Silakan login terlebih dahulu untuk mengakses halaman ini", "warning")
        return redirect(url_for('landingpage'))
    username = session.get('username') or session.get('admin_username', 'Unknown User')
    role = 'user' if 'username' in session else 'admin'
    log_activity(username, role, "Membuka FAQ ")
    return render_template('faq.html')

# Route untuk Tentang Aplikasi
@app.route('/about_app')
def about_app():
    if 'username' not in session and 'admin_username' not in session:
        flash("Silakan login terlebih dahulu untuk mengakses halaman ini", "warning")
        return redirect(url_for('landingpage'))
    username = session.get('username') or session.get('admin_username', 'Unknown User')
    role = 'user' if 'username' in session else 'admin'
    log_activity(username, role, "Membuka Tentang Aplikasi ")
    return render_template('about_app.html')

@app.route('/settings', methods=['GET', 'POST'])
def settings():
    if 'username' not in session and 'admin_username' not in session:
        flash("Silakan login terlebih dahulu untuk mengakses halaman ini", "warning")
        return redirect(url_for('landingpage'))
    username = session.get('username') or session.get('admin_username', 'Unknown User')
    role = 'user' if 'username' in session else 'admin'
    username = session.get('username')
    admin_username = session.get('admin_username')

    if not username and not admin_username:
        flash('Silakan login terlebih dahulu.', 'warning')
        return redirect(url_for('landingpage'))

    success = False

    if request.method == 'POST':
        old_password = request.form['old_password']
        new_password = request.form['new_password']
        confirm_new_password = request.form['confirm_new_password']

        if new_password != confirm_new_password:
            flash('Password baru dan konfirmasi password tidak cocok.', 'error')
            return redirect(url_for('settings'))

        conn = get_db_connection()
        cursor = conn.cursor(dictionary=True)

        if username:
            cursor.execute("SELECT * FROM users WHERE username = %s", (username,))
            role = 'user'
            current_username = username
        else:
            cursor.execute("SELECT * FROM admins WHERE username = %s", (admin_username,))
            role = 'admin'
            current_username = admin_username

        user_or_admin = cursor.fetchone()

        if user_or_admin:
            stored_password_hash = user_or_admin['password']
            print(f"Hash password tersimpan di database: {stored_password_hash}")  # Debugging hash password dari DB
            print(f"Password lama input user: {old_password}")  # Debugging password lama yang diinput

            if check_password_hash(stored_password_hash, old_password):
                print("Password cocok")  # Debugging jika password cocok
                hashed_new_password = generate_password_hash(new_password)
                if role == 'user':
                    cursor.execute("UPDATE users SET password = %s WHERE username = %s", (hashed_new_password, username))
                else:
                    cursor.execute("UPDATE admins SET password = %s WHERE username = %s", (hashed_new_password, admin_username))
                conn.commit()
                logout_time = datetime.now()  # Make sure to import datetime at the top
                cursor.execute(
                    "UPDATE login_records SET logout_time = %s WHERE username = %s AND role = %s AND logout_time IS NULL",
                    (logout_time, current_username, role)
                )
                conn.commit()

                log_activity(current_username, role, 'Change Password')
                flash('Password berhasil diubah! Silakan login kembali.', 'success')
                success = True

                log_activity(current_username, role, 'Change Password')
                flash('Password berhasil diubah!', 'success')
                success = True
            else:
                flash('Password lama tidak sesuai.', 'error')
                print("Password lama tidak cocok")  # Debugging jika password lama tidak cocok
        else:
            print("Admin atau user tidak ditemukan.")


        cursor.close()
        conn.close()

        if success:
            return render_template('settings.html', success=True)
        else:
            return redirect(url_for('settings'))

    return render_template('settings.html')


import random
import string

# Fungsi untuk menghasilkan token reset password
def generate_reset_token(email):
    serializer = URLSafeTimedSerializer(app.config['SECRET_KEY'])
    return serializer.dumps(email, salt='password-reset-salt')

# Fungsi untuk memverifikasi token reset password
def verify_reset_token(token, expiration=3600):
    serializer = URLSafeTimedSerializer(app.config['SECRET_KEY'])
    try:
        email = serializer.loads(token, salt='password-reset-salt', max_age=expiration)
    except:
        return None
    return email
def generate_token(length=50):
    return ''.join(random.choices(string.ascii_letters + string.digits, k=length))



# Route untuk halaman lupa kata sandi
@app.route('/forgot_password', methods=['GET', 'POST'])
def forgot_password():
    if request.method == 'POST':
        email = request.form['email']

        # Cek apakah email ada di database (sesuaikan dengan struktur database Anda)
        conn = get_db_connection()
        cursor = conn.cursor(dictionary=True)
        cursor.execute("SELECT * FROM users WHERE email = %s", (email,))
        user = cursor.fetchone()

        if user:
            # Generate token untuk reset password
            token = generate_token()
            reset_url = url_for('reset_password', token=token, _external=True)

            # Simpan token ke database (untuk verifikasi nanti)
            cursor.execute("UPDATE users SET reset_token = %s WHERE email = %s", (token, email))
            conn.commit()

            # Kirim email dengan link reset password menggunakan Elastic Email
            sender_email = "raflijrockdewasg@gmail.com"  # Gunakan email Elastic Email Anda
            sender_password = "30F58C3EFDFED0280B0ED3459F54C6ED31EF"  # SMTP password dari Elastic Email
            subject = "Reset Password"
            body = f"Hi {user['username']},\n\nKlik link di bawah ini untuk mereset kata sandi Anda:\n{reset_url}\n\nJika Anda tidak meminta reset kata sandi, abaikan email ini."

            # Mempersiapkan email
            msg = MIMEMultipart()
            msg['From'] = sender_email
            msg['To'] = email
            msg['Subject'] = subject
            msg.attach(MIMEText(body, 'plain'))

            try:
                # Mengirim email melalui server SMTP Elastic Email
                with smtplib.SMTP('smtp.elasticemail.com', 2525) as server:
                    server.starttls()  # Menggunakan TLS untuk keamanan
                    server.login(sender_email, sender_password)
                    server.sendmail(sender_email, email, msg.as_string())
                    flash('Link reset kata sandi telah dikirim ke email Anda.', 'info')
            except Exception as e:
                flash(f'Terjadi kesalahan dalam mengirim email: {str(e)}', 'error')
        else:
            flash('Email tidak ditemukan.', 'error')

        cursor.close()
        conn.close()

    return render_template('forgot_password.html')


# Route untuk halaman reset kata sandi
@app.route('/reset_password/<token>', methods=['GET', 'POST'])
def reset_password(token):
    if 'username' not in session and 'admin_username' not in session:
        flash("Silakan login terlebih dahulu untuk mengakses halaman ini", "warning")
        return redirect(url_for('landingpage'))
    username = session.get('username') or session.get('admin_username', 'Unknown User')
    role = 'user' if 'username' in session else 'admin'
    log_activity(username, role, "Mereset Password")
    if request.method == 'POST':
        new_password = request.form['new_password']
        confirm_password = request.form['confirm_password']

        if new_password != confirm_password:
            flash('Password baru dan konfirmasi tidak cocok.', 'error')
            return redirect(url_for('reset_password', token=token))

        # Cari user berdasarkan token
        conn = get_db_connection()
        cursor = conn.cursor(dictionary=True)
        cursor.execute("SELECT * FROM users WHERE reset_token = %s", (token,))
        user = cursor.fetchone()

        if user:
            hashed_password = generate_password_hash(new_password)
            cursor.execute("UPDATE users SET password = %s, reset_token = NULL WHERE reset_token = %s", (hashed_password, token))
            conn.commit()

            flash('Password berhasil diubah! Silakan login dengan password baru Anda.', 'success')
            return redirect(url_for('login'))
        else:
            flash('Token reset tidak valid.', 'error')

        cursor.close()
        conn.close()

    return render_template('reset_password.html')

def get_all_admins():
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SELECT username FROM admins")
    admins = cursor.fetchall()
    cursor.close()
    conn.close()
    return admins

def get_all_users():
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SELECT username, email, nomorhp FROM users")
    users = cursor.fetchall()
    cursor.close()
    conn.close()
    return users

def add_admin(username, password):
    hashed_password = generate_password_hash(password)
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("INSERT INTO admins (username, password) VALUES (%s, %s)", (username, hashed_password))
    conn.commit()
    cursor.close()
    conn.close()

def delete_admin(username):
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM admins WHERE username = %s", (username,))
    conn.commit()
    cursor.close()
    conn.close()

def delete_user(username):
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM users WHERE username = %s", (username,))
    conn.commit()
    cursor.close()
    conn.close()

def add_user(username, email, password, nomorhp):
    hashed_password = generate_password_hash(password)
    conn = get_db_connection()
    cursor = conn.cursor()
    query = "INSERT INTO users (username, email, password, nomorhp) VALUES (%s, %s, %s, %s)"
    cursor.execute(query, (username, email, hashed_password, nomorhp))
    conn.commit()
    cursor.close()
    conn.close()

# Rute utama untuk halaman beranda
@app.route('/home_analisis')
def home_analisis():
    return render_template('home_analisis.html')

# Rute utama untuk halaman cdd/edd
@app.route('/hal_cdd')
def hal_cdd():
    return render_template('cdd.html')

# Rute utama untuk halaman risiko_tppu
@app.route('/hal_risikotppu')
def hal_risikotppu():
    return render_template('risiko_tppu.html')


# Rute utama untuk halaman ira
@app.route('/hal_ira')
def hal_ira():
    return render_template('ira.html')

# Rute utama untuk halaman rbs
@app.route('/hal_rbs')
def hal_rbs():
    return render_template('rbs.html')

# Rute utama untuk halaman sipesat
@app.route('/hal_sipesat')
def hal_sipesat():
    return render_template('sipesat.html')
from fuzzywuzzy import fuzz
from fuzzywuzzy import process

@app.route('/ira', methods=['POST'])
def ira():
    if 'username' not in session and 'admin_username' not in session:
        flash("Silakan login terlebih dahulu untuk mengakses halaman ini", "warning")
        return redirect(url_for('landingpage'))

    username = session.get('username') or session.get('admin_username', 'Unknown User')
    role = 'user' if 'username' in session else 'admin'
    log_activity(username, role, "Membuat Laporan IRA")
    global processed_df, finalized_df

    if request.method == 'POST':
        # Upload Excel files
        file1 = request.files['file1']
        file2 = request.files['file2']
        file3 = request.files['file3']

      # Membuka template laporan Excel
      # Pastikan file ada dan dapat dibaca
    if not file1 or not file2 or not file3:
          return jsonify({"error": "Missing one or more files"}), 400

      # Cek jika file berada di folder yang benar
    current_dir = os.path.dirname(os.path.abspath(__file__))
    template_path = os.path.join(current_dir, 'template_laporan_ira.xlsx')

    if not os.path.exists(template_path):
            return jsonify({"error": f"Template file not found: {template_path}"}), 400

    wb = load_workbook(template_path)
    ws = wb.active

          # Membaca file Excel ke dalam DataFrame
    df1 = pd.read_excel(file1)
    df2 = pd.read_excel(file2)
    df3 = pd.read_excel(file3)

      # (Logika perhitungan seperti sebelumnya)
      # Menghitung jumlah data per 'Fund Name'
    count_fund = df1['Fund Name'].value_counts()

      # Menghitung jumlah data per 'Fund detail name'
    count_fund_detail = df1['Fund Name'].value_counts()

      # Menghitung jumlah data per 'Occupation'
    count_occupation = df2['Occupation'].value_counts()

    count_provinsi = df2["Correspondence City Name"].value_counts() & df3["Company City Name"].value_counts()

      # Menghitung jumlah data per 'Company Type'
    count_company_type = df3['Company Type'].value_counts()

      # Mapping Occupation
    occupation_mapping = {
          'Pejabat Negara/PEP': ['Pejabat Negara', 'Politically Exposed Person', 'PEP'],
          'Karyawan Swasta': ['Karyawan Swasta', 'Private Employee'],
          'Ibu Rumah Tangga': ['Ibu Rumah Tangga', 'Housewife'],
          'Pelajar / Mahasiswa': ['Pelajar', 'Mahasiswa', 'Student', "Student/Graduate Student/Post Graduate"],
          'Polisi Republik Indonesia': ['Polisi Republik Indonesia', 'Police Officer','Indonesian National Armed Force/Indonesian National Police'],
          'Pegawai Negeri Sipil (PNS)': ['Pegawai Negeri Sipil', 'Civil Servant', 'PNS'],
          'Pengurus dan Pegawai BUMN': ['Pengurus BUMN', 'Pegawai BUMN', 'State-Owned Enterprise Employee'],
          'Pengurus dan Pegawai BUMD': ['Pengurus BUMD', 'Pegawai BUMD', 'Regional-Owned Enterprise Employee'],
          'Wirausaha / Wiraswasta': ['Wirausaha', 'Wiraswasta', 'Enterpreneur'],
          'Lainnya ' : ['Others','Other','lainnya'],
          'Tentara Nasional Indonesia (TNI)': ['Tentara Nasional Indonesia', 'Indonesian National Army', 'TNI','Indonesian National Armed Force/Indonesian National Police'],
          'Profesional': ['Profesional', 'Professional'],
          'Pengurus dan Pegawai Yayasan atau LBH': ['Pengurus Yayasan', 'Pegawai Yayasan', 'LBH Staff', 'Foundation or Legal Aid Bureau Staff'],
          'Pengurus dan Pegawai LSM': ['Pengurus LSM', 'Pegawai LSM', 'NGO Staff'],
          'Pensiunan': ['Pensiunan', 'Retiree', 'Retirement'],
          'Pengajar': ['Pengajar', 'Teacher', 'Educator', 'Lecturer', 'Lecturer/Teacher'],
          'Pemuka Agama': ['Pemuka Agama', 'Religious Leader'],
          'Artis/Youtuber/Selebgram/Influencer': ['Artis', 'Youtuber', 'Selebgram', 'Influencer', 'Artist', 'Content Creator'],
          'Sopir': ['Sopir', 'Driver'],
          'Asisten Rumah Tangga': ['Asisten Rumah Tangga', 'Household Assistant', 'Domestic Worker'],
          'Buruh': ['Buruh', 'Laborer'],
          'Atlet/Olahragawan': ['Atlet', 'Olahragawan', 'Athlete', 'Sportsperson'],
          'Tenaga Keamanan': ['Tenaga Keamanan', 'Security Guard']
      }

      # Menempatkan hasil ke dalam template
    start_row_occupation = 31
    for i, (occupation, eng_types) in enumerate(occupation_mapping.items()):
          total_count = sum(count_occupation.get(eng_type, 0) for eng_type in eng_types)
          ws[f"B{start_row_occupation + i}"] = total_count if total_count > 0 else '-'

      # Map untuk menempatkan data ke dalam template sesuai kategori yang ada
    company_type_mapping = {
                  'Partai Politik': ['Partai Politik', 'Political Party'],
                  'Perseroan Terbatas': ['Perseroan Terbatas', 'Limited Liability Company','Corporate','Securites Company','Financial Institution'],
                  'Lembaga Pendidikan Yayasan': ['Lembaga Pendidikan Yayasan', 'Educational Foundation'],
                  'Assuransi': ['Insurance','Asuransi'],
                  'Dana Pensiun': ['Dana Pensiun', 'Pension Fund'],
                  'Koperasi': ['Koperasi', 'Cooperative'],
                  'Organisasi Kemasyarakatan berbentuk yayasan': ['Organisasi Kemasyarakatan berbentuk yayasan', 'Community Organization in the form of Foundation','Foundation']
                  # ada data Lainnya yaitu 'Others','Other' yang tidak tau dikelompokkan kemana
      }

      # Menempatkan hasil perhitungan ke dalam template di baris 56 - 62 dan kolom tertentu
    start_row = 56
    for i, (company, eng_types) in enumerate(company_type_mapping.items()):
          total_count = 0
          for eng_type in eng_types:
              total_count += count_company_type.get(eng_type, 0)

          # Menentukan kolom untuk setiap baris sesuai permintaan
          if i == 0:  # Baris 56 diisi di kolom Tinggi
              ws[f"D{start_row + i}"] = total_count if total_count > 0 else '-'
              ws[f"G{start_row + i}"] = 300 * total_count if total_count > 0 else '-'
          elif i == 1:  # Baris 57 diisi di kolom Sedang
              ws[f"C{start_row + i}"] = total_count if total_count > 0 else '-'
              ws[f"F{start_row + i}"] = 200 * total_count if total_count > 0 else '-'
          else:  # Baris 58-62 diisi di kolom Rendah
              ws[f"B{start_row + i}"] = total_count if total_count > 0 else '-'
              ws[f"E{start_row + i}"] = 100 * total_count if total_count > 0 else '-'

      # Menambahkan kategori untuk Bidang Usaha Nasabah Korporasi di baris 65 - 71
    business_sector_mapping = {
          'Aktivitas Keuangan dan Asuransi': ['Insurance', 'Financial Institution'],  # Asuransi dan Lembaga Keuangan
          'Aktivitas Jasa Lainnya': ['Corporate', 'Others', 'Securities company', 'Pension fund', 'Cooperative',
                                    'Educational foundation', 'Community organization'],
          'Pertambangan dan Penggalian' : ['Mining','Quary','Oil Mining','Gold Mining','Mineral Mining'],
          'Real Estat' : ['Real Estate','House Agency'],
          'Kesenian, Hiburan, dan Rekreasi' : ['Entertaiment','Tv Station','Broadcast Company'],
          'Organisasi Kemasyarakatan berbentuk yayasan' : ['Organisasi Kemasyarakatan berbentuk yayasan', 'Community Organization in the form of Foundation','Foundation']
      }

      # Menempatkan hasil perhitungan untuk Bidang Usaha Nasabah Korporasi di baris 65 - 71
    start_row_business = 65
    for i, (sector, company_types) in enumerate(business_sector_mapping.items()):
          total_count = 0
          for company_type in company_types:
              total_count += count_company_type.get(company_type, 0)

          # Menentukan kolom untuk setiap baris (Rendah, Sedang, Tinggi)
          if i == 0:  # Aktifitas Keuangan dan Asuransi diisi di kolom Tinggi
              ws[f"C{start_row_business + i}"] = total_count if total_count > 0 else '-'
              ws[f"F{start_row_business + i}"] = 200 * total_count if total_count > 0 else '-'
          elif i == 1:  # Aktifitas Jasa Lainnya diisi di kolom Sedang
              ws[f"C{start_row_business + i}"] = total_count if total_count > 0 else '-'
              ws[f"F{start_row_business + i}"] = 200 * total_count if total_count > 0 else '-'
          else:  # Baris 70-71 diisi di kolom Rendah
              ws[f"B{start_row_business + i}"] = total_count if total_count > 0 else '-'
              ws[f"E{start_row_business + i}"] = 100 * total_count if total_count > 0 else '-'

      # Daftar kategori fund
    fund_mapping = {
          "REKSA DANA SAHAM" : ["PNM SAHAM AGRESIF", "REKSA DANA PNM EKUITAS SYARIAH","REKSA DANA PNM SAHAM UNGGULAN"],
          "REKSA DANA INDEKS" : ["REKSA DANA INDEKS PNM INDEKS INFOBANK15"],
          "REKSA DANA CAMPURAN" : ["REKSA DANA PNM SYARIAH"],
          "REKSA DANA PENDAPATAN TETAP" : [ "REKSA DANA PNM AMANAH SYARIAH KELAS A", "REKSADANA PNM DANA SEJAHTERA II",
                                               "REKSA DANA PNM DANA BERTUMBUH", "REKSA DANA PNM DANA SURAT BERHARGA NEGARA",
                                               "REKSA DANA PNM DANA SURAT BERHARGA NEGARA II KELAS A", "REKSADANA SYARIAH PNM SUKUK NEGARA SYARIAH",
                                               "Reksa Dana PNM SBN 90", "Reksa Dana Syariah Pendapatan Tetap PNM Kaffah",
                                               "REKSA DANA PENDAPATAN TETAP PNM DANA OPTIMA KELAS A",
                                               "REKSA DANA SYARIAH PENDAPATAN TETAP PNM SURAT BERHARGA SYARIAH NEGARA",
                                               "REKSA DANA PENDAPATAN TETAP PNM OPTIMA BULANAN"],
          "KONTRAK PENGELOLAAN DANA (KPD)" : ["KONTRAK PENGELOLAAN DANA (KPD)"],
          "REKSA DANA PASAR UANG" :[ "REKSA DANA PNM PUAS", "REKSA DANA PNM DANA TUNAI", "REKSA DANA PNM PASAR UANG SYARIAH",
                                    "REKSA DANA PNM DANA KAS PLATINUM", "REKSA DANA PNM DANA LIKUID",
                                    "REKSA DANA SYARIAH PASAR UANG PNM ARAFAH", "REKSA DANA SYARIAH PASAR UANG PNM FALAH",
                                    "REKSA DANA SYARIAH PASAR UANG PNM FALAH 2", "REKSA DANA SYARIAH PASAR UANG PNM FALAH 3",
                                    "REKSA DANA SYARIAH PASAR UANG PNM FAAZA", "REKSA DANA PASAR UANG PNM DANA KAS PLATINUM 2",
                                    "REKSA DANA PASAR UANG PNM DANA MAXIMA", "REKSA DANA PASAR UANG PNM DANA MAXIMA 2"],
          "REKSA DANA TERPROTEKSI" : ["REKSA DANA SYARIAH TERPROTEKSI PNM TERPROTEKSI INVESTA 40",
                                      "REKSA DANA SYARIAH TERPROTEKSI PNM TERPROTEKSI INVESTA 44",
                                      "REKSA DANA TERPROTEKSI PNM TERPROTEKSI INVESTA 41", "REKSA DANA TERPROTEKSI PNM TERPROTEKSI INVESTA 42"],
          "REKSA DANA PENYERTAAN TERBATAS (RDPT)" : ["REKSA DANA PENYERTAAN TERBATAS PNM ADHI GUNA PUTERA",
                                                     "REKSA DANA PENYERTAAN TERBATAS PNM PERIKANAN NUSANTARA", "REKSA DANA PENYERTAAN TERBATAS PNM VENTURE CAPITAL",
                                                     "REKSA DANA SYARIAH PENYERTAAN TERBATAS PNM INDAH KARYA",
                                                     "REKSA DANA SYARIAH PENYERTAAN TERBATAS PNM PEMBIAYAAN MIKRO BUMN SERI IX",
                                                     "REKSA DANA SYARIAH PENYERTAAN TERBATAS PNM PEMBIAYAAN MIKRO BUMN SERI VI",
                                                     "REKSA DANA SYARIAH PENYERTAAN TERBATAS PNM PEMBIAYAAN MIKRO BUMN SERI VII",
                                                     "REKSA DANA SYARIAH PENYERTAAN TERBATAS PNM PEMBIAYAAN MIKRO BUMN SERI XI"],
         #tidak terdapat data ini
         "DANA INVESTASI REAL ESTAT (DIRE)" : ["DANA INVESTASI REAL ESTAT (DIRE)"],
         "DANA INVESTASI INFRASTRUKTUR (DINFRA)" : ["DANA INVESTASI INFRASTRUKTUR (DINFRA)"],
         "EFEK BERAGUN ASET (EBA)" : ["EFEK BERAGUN ASET (EBA)"]
          }

      # Menempatkan hasil perhitungan untuk PRODUK/JASA di baris 74 - 84
    start_row_fund = 74
    for i, (fund, fund_types) in enumerate(fund_mapping.items()):
          total_count = 0
          for fund_type in fund_types:
              total_count += count_fund.get(fund_type, 0)

          # Menentukan kolom untuk setiap baris (Rendah, Sedang, Tinggi)
          if i == 0:  # REKSADANA SAHAM DAN INDEKS SEDANG
              ws[f"C{start_row_fund + i}"] = total_count if total_count > 0 else '-'
              ws[f"F{start_row_fund + i}"] = 200 * total_count if total_count > 0 else '-'
          elif i == 1:  # Aktifitas Jasa Lainnya diisi di kolom Sedang
              ws[f"C{start_row_fund + i}"] = total_count if total_count > 0 else '-'
              ws[f"F{start_row_fund + i}"] = 200 * total_count if total_count > 0 else '-'
          else:  # Baris 70-71 diisi di kolom Rendah
              ws[f"B{start_row_fund + i}"] = total_count if total_count > 0 else '-'
              ws[f"E{start_row_fund + i}"] = 100 * total_count if total_count > 0 else '-'

      # Daftar kategori fund detail
    fund_detail_mapping = {
          #REKSA DANA SAHAM
          "REKSA DANA SAHAM" : ["PNM SAHAM AGRESIF", "REKSA DANA PNM EKUITAS SYARIAH","REKSA DANA PNM SAHAM UNGGULAN"],
          "PNM SAHAM AGRESIF" : ["PNM SAHAM AGRESIF"],
          "REKSA DANA PNM EKUITAS SYARIAH" : ["REKSA DANA PNM EKUITAS SYARIAH"],
          "REKSA DANA PNM SAHAM UNGGULAN" : ["REKSA DANA PNM SAHAM UNGGULAN"],

          #REKSA DANA INDEKS
          "REKSA DANA INDEKS" : ["REKSA DANA INDEKS PNM INDEKS INFOBANK15"],
          "REKSA DANA INDEKS PNM INDEKS INFOBANK15" : ["REKSA DANA INDEKS PNM INDEKS INFOBANK15"],

          #REKSA DANA CAMPURAN
          "REKSA DANA CAMPURAN" : ["REKSA DANA PNM SYARIAH"],
          "REKSA DANA PNM SYARIAH" : ["REKSA DANA PNM SYARIAH"],

          #REKSA DANA PENDAPATAN TETAP
          "REKSA DANA PENDAPATAN TETAP" : [ "REKSA DANA PNM AMANAH SYARIAH KELAS A", "REKSADANA PNM DANA SEJAHTERA II",
                                               "REKSA DANA PNM DANA BERTUMBUH", "REKSA DANA PNM DANA SURAT BERHARGA NEGARA",
                                               "REKSA DANA PNM DANA SURAT BERHARGA NEGARA II KELAS A", "REKSADANA SYARIAH PNM SUKUK NEGARA SYARIAH",
                                               "Reksa Dana PNM SBN 90", "Reksa Dana Syariah Pendapatan Tetap PNM Kaffah",
                                               "REKSA DANA PENDAPATAN TETAP PNM DANA OPTIMA KELAS A",
                                               "REKSA DANA SYARIAH PENDAPATAN TETAP PNM SURAT BERHARGA SYARIAH NEGARA",
                                               "REKSA DANA PENDAPATAN TETAP PNM OPTIMA BULANAN"],
          "REKSA DANA PNM AMANAH SYARIAH KELAS A" : ["REKSA DANA PNM AMANAH SYARIAH KELAS A"],
          "REKSADANA PNM DANA SEJAHTERA II" : ["REKSADANA PNM DANA SEJAHTERA II"],
          "REKSA DANA PNM DANA BERTUMBUH" : ["REKSA DANA PNM DANA BERTUMBUH"],
          "REKSA DANA PNM DANA SURAT BERHARGA NEGARA" : ["REKSA DANA PNM DANA SURAT BERHARGA NEGARA"],
          "REKSA DANA PNM DANA SURAT BERHARGA NEGARA II KELAS A" : ["REKSA DANA PNM DANA SURAT BERHARGA NEGARA II KELAS A"],
          "REKSADANA SYARIAH PNM SUKUK NEGARA SYARIAH" : ["REKSADANA SYARIAH PNM SUKUK NEGARA SYARIAH"],
          "Reksa Dana PNM SBN 90" : ["Reksa Dana PNM SBN 90"],
          "Reksa Dana Syariah Pendapatan Tetap PNM Kaffah" : ["Reksa Dana Syariah Pendapatan Tetap PNM Kaffah"],
          "REKSA DANA PENDAPATAN TETAP PNM DANA OPTIMA KELAS A" : ["REKSA DANA PENDAPATAN TETAP PNM DANA OPTIMA KELAS A"],
          "REKSA DANA SYARIAH PENDAPATAN TETAP PNM SURAT BERHARGA SYARIAH NEGARA" : ["REKSA DANA SYARIAH PENDAPATAN TETAP PNM SURAT BERHARGA SYARIAH NEGARA"],
          "REKSA DANA PENDAPATAN TETAP PNM OPTIMA BULANAN" : ["REKSA DANA PENDAPATAN TETAP PNM OPTIMA BULANAN"],

          #REKSA DANA PASAR UANG
          "REKSA DANA PASAR UANG" :[ "REKSA DANA PNM PUAS", "REKSA DANA PNM DANA TUNAI", "REKSA DANA PNM PASAR UANG SYARIAH",
                                    "REKSA DANA PNM DANA KAS PLATINUM", "REKSA DANA PNM DANA LIKUID",
                                    "REKSA DANA SYARIAH PASAR UANG PNM ARAFAH", "REKSA DANA SYARIAH PASAR UANG PNM FALAH",
                                    "REKSA DANA SYARIAH PASAR UANG PNM FALAH 2", "REKSA DANA SYARIAH PASAR UANG PNM FALAH 3",
                                    "REKSA DANA SYARIAH PASAR UANG PNM FAAZA", "REKSA DANA PASAR UANG PNM DANA KAS PLATINUM 2",
                                    "REKSA DANA PASAR UANG PNM DANA MAXIMA", "REKSA DANA PASAR UANG PNM DANA MAXIMA 2"],
          "REKSA DANA PNM PUAS": ["REKSA DANA PNM PUAS"],
          "REKSA DANA PNM DANA TUNAI" : ["REKSA DANA PNM DANA TUNAI"],
          "REKSA DANA PNM PASAR UANG SYARIAH" : ["REKSA DANA PNM PASAR UANG SYARIAH"],
          "REKSA DANA PNM DANA KAS PLATINUM" : ["REKSA DANA PNM DANA KAS PLATINUM"],
          "REKSA DANA PNM DANA LIKUID" : ["REKSA DANA PNM DANA LIKUID"],
          "REKSA DANA SYARIAH PASAR UANG PNM ARAFAH" : ["REKSA DANA SYARIAH PASAR UANG PNM ARAFAH","Reksa Dana Syariah Pasar Uang PNM Arafah"],
          "REKSA DANA SYARIAH PASAR UANG PNM FALAH" : ["REKSA DANA SYARIAH PASAR UANG PNM FALAH"],
          "REKSA DANA SYARIAH PASAR UANG PNM FALAH 2" : ["REKSA DANA SYARIAH PASAR UANG PNM FALAH 2"],
          "REKSA DANA SYARIAH PASAR UANG PNM FALAH 3" : ["REKSA DANA SYARIAH PASAR UANG PNM FALAH 3"],
          "REKSA DANA SYARIAH PASAR UANG PNM FAAZA" : ["REKSA DANA SYARIAH PASAR UANG PNM FAAZA"],
          "REKSA DANA PASAR UANG PNM DANA KAS PLATINUM 2" : ["REKSA DANA PASAR UANG PNM DANA KAS PLATINUM 2"],
          "REKSA DANA PASAR UANG PNM DANA MAXIMA" : ["REKSA DANA PASAR UANG PNM DANA MAXIMA"],
          "REKSA DANA PASAR UANG PNM DANA MAXIMA 2" : ["REKSA DANA PASAR UANG PNM DANA MAXIMA 2"],

          #REKSA DANA TERPROTEKSI
          "REKSA DANA TERPROTEKSI" : ["REKSA DANA SYARIAH TERPROTEKSI PNM TERPROTEKSI INVESTA 40",
                                      "REKSA DANA SYARIAH TERPROTEKSI PNM TERPROTEKSI INVESTA 44",
                                      "REKSA DANA TERPROTEKSI PNM TERPROTEKSI INVESTA 41", "REKSA DANA TERPROTEKSI PNM TERPROTEKSI INVESTA 42"],
          "REKSA DANA SYARIAH TERPROTEKSI PNM TERPROTEKSI INVESTA 40" : ["REKSA DANA SYARIAH TERPROTEKSI PNM TERPROTEKSI INVESTA 40"],
          "REKSA DANA SYARIAH TERPROTEKSI PNM TERPROTEKSI INVESTA 44" : ["REKSA DANA SYARIAH TERPROTEKSI PNM TERPROTEKSI INVESTA 44"],
          "REKSA DANA TERPROTEKSI PNM TERPROTEKSI INVESTA 41" : ["REKSA DANA TERPROTEKSI PNM TERPROTEKSI INVESTA 41"],
          "REKSA DANA TERPROTEKSI PNM TERPROTEKSI INVESTA 42" : ["REKSA DANA TERPROTEKSI PNM TERPROTEKSI INVESTA 42"],

          #REKSA DANA PENYERTAAN TERBATAS (RDPT)
          "REKSA DANA PENYERTAAN TERBATAS (RDPT)" : ["REKSA DANA PENYERTAAN TERBATAS PNM ADHI GUNA PUTERA",
                                                     "REKSA DANA PENYERTAAN TERBATAS PNM PERIKANAN NUSANTARA", "REKSA DANA PENYERTAAN TERBATAS PNM VENTURE CAPITAL",
                                                     "REKSA DANA SYARIAH PENYERTAAN TERBATAS PNM INDAH KARYA",
                                                     "REKSA DANA SYARIAH PENYERTAAN TERBATAS PNM PEMBIAYAAN MIKRO BUMN SERI IX",
                                                     "REKSA DANA SYARIAH PENYERTAAN TERBATAS PNM PEMBIAYAAN MIKRO BUMN SERI VI",
                                                     "REKSA DANA SYARIAH PENYERTAAN TERBATAS PNM PEMBIAYAAN MIKRO BUMN SERI VII",
                                                     "REKSA DANA SYARIAH PENYERTAAN TERBATAS PNM PEMBIAYAAN MIKRO BUMN SERI XI"],
          "REKSA DANA PENYERTAAN TERBATAS PNM ADHI GUNA PUTERA" : ["REKSA DANA PENYERTAAN TERBATAS PNM ADHI GUNA PUTERA"],
          "REKSA DANA PENYERTAAN TERBATAS PNM PERIKANAN NUSANTARA" : ["REKSA DANA PENYERTAAN TERBATAS PNM PERIKANAN NUSANTARA"],
          "REKSA DANA PENYERTAAN TERBATAS PNM VENTURE CAPITAL" : ["REKSA DANA PENYERTAAN TERBATAS PNM VENTURE CAPITAL"],
          "REKSA DANA SYARIAH PENYERTAAN TERBATAS PNM INDAH KARYA" : ["REKSA DANA SYARIAH PENYERTAAN TERBATAS PNM INDAH KARYA"],
          "REKSA DANA SYARIAH PENYERTAAN TERBATAS PNM PEMBIAYAAN MIKRO BUMN SERI IX" : ["REKSA DANA SYARIAH PENYERTAAN TERBATAS PNM PEMBIAYAAN MIKRO BUMN SERI IX"],
          "REKSA DANA SYARIAH PENYERTAAN TERBATAS PNM PEMBIAYAAN MIKRO BUMN SERI VI" : ["REKSA DANA SYARIAH PENYERTAAN TERBATAS PNM PEMBIAYAAN MIKRO BUMN SERI VI"],
          "REKSA DANA SYARIAH PENYERTAAN TERBATAS PNM PEMBIAYAAN MIKRO BUMN SERI VII" : ["REKSA DANA SYARIAH PENYERTAAN TERBATAS PNM PEMBIAYAAN MIKRO BUMN SERI VII"],
          "REKSA DANA SYARIAH PENYERTAAN TERBATAS PNM PEMBIAYAAN MIKRO BUMN SERI XI" : ["REKSA DANA SYARIAH PENYERTAAN TERBATAS PNM PEMBIAYAAN MIKRO BUMN SERI XI"],

          #KONTRAK PENGELOLAAN DANA (KPD)
          "KONTRAK PENGELOLAAN DANA (KPD)" : ["KONTRAK PENGELOLAAN DANA (KPD)"],
          "KONTRAK PENGELOLAAN DANA (KPD)" : ["KONTRAK PENGELOLAAN DANA (KPD)"]
         #sengaja di nonaktifkan karena tidak terdapat bentuk data tersebut
         #"DANA INVESTASI REAL ESTAT (DIRE)" : ["DANA INVESTASI REAL ESTAT (DIRE)"],
         #"DANA INVESTASI INFRASTRUKTUR (DINFRA)" : ["DANA INVESTASI INFRASTRUKTUR (DINFRA)"],
         #"EFEK BERAGUN ASET (EBA)" : ["EFEK BERAGUN ASET (EBA)"]
          }
      # Menempatkan hasil perhitungan untuk PRODUK/JASA di baris 86 - 135
    start_row_fund_detail = 86
    for i, (fund, fund_detail_types) in enumerate(fund_detail_mapping.items()):
          total_count = 0
          for fund_detail_type in fund_detail_types:
              total_count += count_fund_detail.get(fund_detail_type, 0)

          # Menentukan kolom untuk setiap baris (Rendah, Sedang, Tinggi)
          if i == 0:  # REKSADANA SAHAM DAN INDEKS SEDANG
              ws[f"C{start_row_fund_detail + i}"] = total_count if total_count > 0 else '-'
              ws[f"F{start_row_fund_detail + i}"] = 200 * total_count if total_count > 0 else '-'
          elif i == 1:  # Aktifitas Jasa Lainnya diisi di kolom Sedang
              ws[f"C{start_row_fund_detail + i}"] = total_count if total_count > 0 else '-'
              ws[f"F{start_row_fund_detail + i}"] = 200 * total_count if total_count > 0 else '-'
          elif i == 2:  # Aktifitas Jasa Lainnya diisi di kolom Sedang
              ws[f"C{start_row_fund_detail + i}"] = total_count if total_count > 0 else '-'
              ws[f"F{start_row_fund_detail + i}"] = 200 * total_count if total_count > 0 else '-'
          elif i == 3:  # Aktifitas Jasa Lainnya diisi di kolom Sedang
              ws[f"C{start_row_fund_detail + i}"] = total_count if total_count > 0 else '-'
              ws[f"F{start_row_fund_detail + i}"] = 200 * total_count if total_count > 0 else '-'
          elif i == 4:  # Aktifitas Jasa Lainnya diisi di kolom Sedang
              ws[f"C{start_row_fund_detail + i}"] = total_count if total_count > 0 else '-'
              ws[f"F{start_row_fund_detail + i}"] = 200 * total_count if total_count > 0 else '-'
          elif i == 5:  # Aktifitas Jasa Lainnya diisi di kolom Sedang
              ws[f"C{start_row_fund_detail + i}"] = total_count if total_count > 0 else '-'
              ws[f"F{start_row_fund_detail + i}"] = 200 * total_count if total_count > 0 else '-'
          else:  # Baris 70-71 diisi di kolom Rendah
              ws[f"B{start_row_fund_detail + i}"] = total_count if total_count > 0 else '-'
              ws[f"E{start_row_fund_detail + i}"] = 100 * total_count if total_count > 0 else '-'

    provinsi_mapping ={
                  "DKI Jakarta": ["Jakarta Pusat", "Jakarta Utara", "Jakarta Barat", "Jakarta Selatan", "Jakarta Timur",
                                  "KODYA JAKARTA PUSAT", "KODYA JAKARTA UTARA", "KODYA JAKARTA BARAT", "KODYA JAKARTA SELATAN", "KODYA JAKARTA TIMUR",
                                  "KOTA JAKARTA PUSAT", "KOTA JAKARTA UTARA", "KOTA JAKARTA BARAT", "KOTA JAKARTA SELATAN", "KOTA JAKARTA TIMUR",
                                  "JAKARTA PUSAT", "JAKARTA UTARA", "JAKARTA BARAT", "JAKARTA SELATAN", "JAKARTA TIMUR","DKI JAKARTA", "JAKARTA"
                                  ],
                  "Jawa Barat": ["Bandung", "Bekasi", "Bogor", "Cimahi", "Depok","Cianjur","Sumedang","Cirebon","Cikarang", "Garut","Sukabumi","Tasikmalaya","Ciamis"],
                  "Jawa Tengah": ["Semarang", "Solo", "Magelang", "Salatiga", "Kudus","Pati", "Tegal","Brebes","Jepara","Sragen","Pekalongan","Karanganyar","Wonosobo","Purwokerto","Wonogiri"],
                  "Jawa Timur": ["Surabaya", "Malang", "Sidoarjo", "Madiun", "Probolinggo","Magetan","Ngawi","Lamongan","Gresik","Tuban","Pasuruan","Jombang","Banyuwangi", "Mojokerto","Kediri","Bojonegoro"],
                  "Sumatera Utara": ["Medan", "Binjai", "Pematang Siantar", "Tanjungbalai", "Sibolga"],
                  "Bali": ["Denpasar", "Badung", "Buleleng", "Gianyar", "Tabanan"],
                  "Sulawesi Tenggara": ["Kendari", "Baubau", "Kolaka", "Muna", "Buton"],
                  "Sulawesi Utara": ["Manado", "Bitung", "Tomohon", "Minahasa", "Sangihe"],
                  "Sulawesi Selatan": ["Makassar", "Pare-pare", "Palopo", "Gowa", "Bulukumba"],
                  "Sulawesi Tengah": ["Palu", "Donggala", "Tolitoli", "Morowali", "Banggai"],
                  "Sulawesi Barat": ["Mamuju", "Majene", "Polman", "Mamasa", "Balanipa"],
                  "Gorontalo": ["Gorontalo", "Bone Bolango", "Pohuwato", "Boalemo", "Gorontalo Utara"],
                  "DI Yogyakarta": ["Yogyakarta", "Sleman", "Bantul", "Kulon Progo", "Gunungkidul"],
                  "Banten": ["Serang", "Tangerang", "Cilegon", "Pandeglang", "Lebak"],
                  "Nangroe Aceh Darussalam": ["Banda Aceh", "Lhokseumawe", "Langsa", "Sabang", "Meulaboh"],
                  "Riau": ["Pekanbaru", "Dumai", "Siak", "Tembilahan", "Bengkalis"],
                  "Sumatera Selatan": ["Palembang", "Prabumulih", "Lubuklinggau", "Baturaja", "Muara Enim"],
                  "Sumatera Barat": ["Padang", "Bukittinggi", "Payakumbuh", "Solok", "Padang Panjang"],
                  "Bangka Belitung": ["Pangkal Pinang", "Toboali", "Muntok", "Sungailiat", "Koba"],
                  "Lampung": ["Bandar Lampung", "Metro", "Pringsewu", "Tulang Bawang", "Lampung Selatan"],
                  "Bengkulu": ["Bengkulu", "Curup", "Ratu Agung", "Mukomuko", "Kepahiang"],
                  "Jambi": ["Jambi", "Sungai Penuh", "Muara Bungo", "Tebo", "Sarolangun"],
                  "Nusa Tenggara Timur": ["Kupang", "Maumere", "Atambua", "Ruteng", "Ende"],
                  "Nusa Tenggara Barat": ["Mataram", "Bima", "Sumbawa", "Dompu", "Lombok"],
                  "Kalimantan Timur": ["Samarinda", "Balikpapan", "Bontang", "Kutai Kartanegara", "Sangatta"],
                  "Kalimantan Barat": ["Pontianak", "Singkawang", "Sintang", "Kubu Raya", "Ketapang"],
                  "Kalimantan Selatan": ["Banjarmasin", "Banjarbaru", "Martapura", "Barabai", "Amuntai"],
                  "Kalimantan Tengah": ["Palangka Raya", "Kuala Pembuang", "Sampit", "Pangkalan Bun", "Tamiang Layang"],
                  "Papua": ["Jayapura", "Manokwari", "Merauke", "Biak", "Timika"],
          }

      # Fungsi untuk mencocokkan nama kota dengan toleransi kesalahan (fuzzy matching)
    def match_kota(kota, kota_list, threshold=10):
          """
          Mencocokkan nama kota dengan daftar kota menggunakan fuzzy matching.
          Jika kecocokan di atas threshold, kembalikan nama kota yang cocok.
          """
          matches = process.extractOne(kota, kota_list, scorer=fuzz.ratio)
          if matches and matches[1] >= threshold:
              return matches[0]
          return None  # Tidak ada kecocokan yang cukup

    start_row_provinsi = 138
    for i, (provinsi, kota_list) in enumerate(provinsi_mapping.items()):
          total_count = 0

          # Menghitung jumlah berdasarkan nama kota dengan fuzzy matching
          for kota in kota_list:
              # Gunakan fuzzy matching untuk mencocokkan nama kota yang lebih fleksibel
              matched_kota = match_kota(kota, count_provinsi.keys())

              # Jika kota yang cocok ditemukan, tambahkan jumlahnya
              if matched_kota:
                  total_count += count_provinsi.get(matched_kota, 0)

          # Menentukan kolom untuk setiap baris (Rendah, Sedang, Tinggi)
          if i == 0:  # REKSADANA SAHAM DAN INDEKS SEDANG
              ws[f"D{start_row_provinsi + i}"] = total_count if total_count > 0 else '-'
              ws[f"G{start_row_provinsi + i}"] = 300 * total_count if total_count > 0 else '-'
          else:  # Baris 70-71 diisi di kolom Rendah
              ws[f"B{start_row_provinsi + i}"] = total_count if total_count > 0 else '-'
              ws[f"E{start_row_provinsi + i}"] = 100 * total_count if total_count > 0 else '-'




  # Fungsi untuk mengisi sel kosong dengan tanda '-'
    def fill_empty_cells(ws, start_row, end_row, columns):
          for row in range(start_row, end_row + 1):
              for col in columns:
                  if ws[f'{col}{row}'].value in [None, ""]:
                      ws[f'{col}{row}'].value = "-"

      # Mengisi sel kosong di rentang tertentu
    ranges = [(4, 28), (31, 53), (56, 62), (65, 71), (74, 84), (86, 135), (138, 167), (170, 174)]
    for start, end in ranges:
          fill_empty_cells(ws, start, end, ['B', 'C', 'D', 'E', 'F', 'G'])

      # Menyimpan laporan ke file baru
      # Variabel tetap dipertahankan, tetapi hanya untuk ilustrasi struktur direktori
    home_dir = os.path.expanduser("~")  # Mendapatkan direktori home user
    downloads_dir = os.path.join(home_dir, "Downloads")  # Mengarah ke folder Downloads
    output_path = os.path.join(downloads_dir, 'output_laporan_ira.xlsx')

    # Menyimpan file ke buffer
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)

    # Return file untuk diunduh
    return send_file(
        output,
        as_attachment=True,
        download_name='output_laporan_ira.xlsx',
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )

    return jsonify({"message": "Laporan berhasil dibuat", "file_path": output_path}), 200
        # Kembalikan atau tampilkan hasil akhir

@app.route('/cdd_edd', methods=['GET', 'POST'])
def cdd_edd():
    if 'username' not in session and 'admin_username' not in session:
        flash("Silakan login terlebih dahulu untuk mengakses halaman ini", "warning")
        return redirect(url_for('landingpage'))

    username = session.get('username') or session.get('admin_username', 'Unknown User')
    role = 'user' if 'username' in session else 'admin'
    log_activity(username, role, "Melakukan Proses CDD & EDD")
    global processed_df, finalized_df

    if request.method == 'POST':
        # Save uploaded files
        file1 = request.files['file1']
        file2 = request.files['file2']
        file3 = request.files['file3']

        # Read Excel files and get relevant columns
        df1 = pd.read_excel(file1, usecols=['SID', 'Investor Fund Unit A/C Name', 'Fund Name', 'CB Name', 'Amount (IDR Equivalent)'])
        df2 = pd.read_excel(file2, usecols=['SID', 'Investor Type', 'First Name', 'Middle Name', 'Last Name', 'ID No.', 'Occupation', 'Income Level (IDR)', 'Asset Owner', 'KTP Address', 'Correspondence Address', 'Country of Domicile'])
        df3 = pd.read_excel(file3, usecols=['SID', 'Investor Type', 'Company Name'])

        # Remove rows without SID from all files
        df1.dropna(subset=['SID'], inplace=True)
        df2.dropna(subset=['SID'], inplace=True)
        df3.dropna(subset=['SID'], inplace=True)

        # Combine customer names into one column
        df2['Name Nasabah'] = df2[['First Name', 'Middle Name', 'Last Name']].fillna('').agg(' '.join, axis=1)

        # Create 'Nama' column for institutional or individual investors
        df3['Nama'] = df3['Company Name']
        df2['Nama'] = df2['Name Nasabah']

        # Filter SIDs that match in the second or third file
        df1_matched_df2 = df1[df1['SID'].isin(df2['SID'])]
        df1_matched_df3 = df1[~df1['SID'].isin(df2['SID']) & df1['SID'].isin(df3['SID'])]

        # Merge dataframes based on matching SIDs
        result_df2 = pd.merge(df1_matched_df2, df2, on='SID', how='left')
        result_df3 = pd.merge(df1_matched_df3, df3[['SID', 'Investor Type', 'Nama']], on='SID', how='left')

        # Concatenate results and aggregate
        final_df = pd.concat([result_df2, result_df3], ignore_index=True)
        final_df = final_df.groupby(['SID', 'Nama', 'Investor Type'], as_index=False).agg({
            'Fund Name': lambda x: ', '.join(x),
            'Amount (IDR Equivalent)': 'sum'
        })

        # Create columns for 'Rendah' and 'Tinggi' risks based on threshold
        final_df['CDD'] = final_df['Amount (IDR Equivalent)'].apply(lambda x: 'X' if x <= 500_000_000 else '')
        final_df['EDD'] = final_df['Amount (IDR Equivalent)'].apply(lambda x: 'X' if x > 500_000_000 else '')

        # Reorder columns
        final_df = final_df[['SID', 'Nama', 'Fund Name', 'Investor Type', 'Amount (IDR Equivalent)', 'CDD', 'EDD']]

        # Rename columns
        final_df.columns = ['SID', 'Nama', 'Reksa Dana', 'Jenis Nasabah', 'Nominal', 'CDD', 'EDD']

        # Sort data based on 'EDD', 'Jenis Nasabah', and 'Nominal' in descending order
        final_df.sort_values(by=['EDD', 'Jenis Nasabah', 'Nominal'], ascending=[False, False, False], inplace=True)

        # Add numbering
        final_df.insert(0, 'No', range(1, len(final_df) + 1))

        # Add a summary row
        summary_row = pd.DataFrame({
            'No': [''],
            'Nama': ['TOTAL DATA'],
            'SID': [''],
            'Reksa Dana': [''],
            'Jenis Nasabah': [''],
            'Nominal': [final_df.shape[0]],
            'CDD': [final_df['CDD'].value_counts().get('X', 0)],
            'EDD': [final_df['EDD'].value_counts().get('X', 0)]
        })

        final_df = pd.concat([final_df, summary_row], ignore_index=True)

        # Store the result for viewing and download
        processed_df = final_df
        finalized_df = processed_df.copy()

        return redirect(url_for('lihat_cdd'))

    return render_template('home_analisis.html')


@app.route('/download_cdd')
def download_cdd():
    if 'username' not in session and 'admin_username' not in session:
        flash("Silakan login terlebih dahulu untuk mengakses halaman ini", "warning")
        return redirect(url_for('landingpage'))

    username = session.get('username') or session.get('admin_username', 'Unknown User')
    role = 'user' if 'username' in session else 'admin'
    log_activity(username, role, "Mengunduh Excel CDD & EDD")
    global finalized_df  # Use the locked DataFrame for download

    if finalized_df is not None:
        workbook = Workbook()
        worksheet = workbook.active

        # Format 'Reksa Dana' for new lines
        finalized_df['Reksa Dana'] = finalized_df['Reksa Dana'].apply(
            lambda x: '\n'.join(f"{i+1}. {name.strip()}" for i, name in enumerate(x.split(',')))
        )

        # Write headers manually
        headers = ['No', 'Nama', 'SID', 'Reksa Dana', 'Jenis Nasabah', 'Nominal']
        worksheet.append(headers)

        # Insert custom header for 'Tingkat Risiko Nasabah'
        worksheet.merge_cells('G1:H1')
        worksheet['G1'] = 'Tingkat Risiko Nasabah'
        worksheet['G1'].alignment = Alignment(horizontal='center', vertical='center')

        # Add subheaders for 'Rendah' and 'Tinggi'
        worksheet['G2'] = 'CDD'
        worksheet['H2'] = 'EDD'

        # Write the DataFrame to the worksheet, starting from row 3
        for r_idx, row in enumerate(dataframe_to_rows(finalized_df, index=False, header=False), 3):
            worksheet.append(row)

        # Apply alignment to header cells
        for col in ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H']:
            worksheet[f'{col}1'].alignment = Alignment(horizontal='center', vertical='center')
            worksheet[f'{col}2'].alignment = Alignment(horizontal='center', vertical='center')

        # Adjust column widths for better readability
        columns = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H']
        for col_letter in columns:
            max_length = 0
            for cell in worksheet[col_letter]:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))

            adjusted_width = max_length + 2
            worksheet.column_dimensions[col_letter].width = adjusted_width

            # Special case for 'Reksa Dana' column (column D), setting a maximum width
            if col_letter == 'D':
                worksheet.column_dimensions[col_letter].width = 80  # Limit the width of the 'Reksa Dana' column

        # Apply text wrapping for 'Reksa Dana' column (Column D)
        for row in worksheet.iter_rows(min_row=3, min_col=4, max_col=4):
            for cell in row:
                cell.alignment = Alignment(wrap_text=True)

        # Apply middle alignment to all data cells
        for row in worksheet.iter_rows(min_row=3):
            for cell in row:
                if cell.column_letter not in ['D']:  # Skip 'Reksa Dana', 'Rendah', and 'Tinggi' columns for center alignment
                    cell.alignment = Alignment(horizontal='center', vertical='center')

        # Save the workbook to a BytesIO object
        output = BytesIO()
        workbook.save(output)
        output.seek(0)

        # Return the Excel file for download
        return send_file(output, as_attachment=True, download_name='result_cdd_edd.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    else:
        return "No file to download"


@app.route('/lihat_cdd')
def lihat_cdd():
    global processed_df

    if processed_df is not None:
        # Format 'Nominal' for better readability
        processed_df['Nominal'] = processed_df['Nominal'].apply(lambda x: f'{x:,.0f}'.replace(',', '.') if isinstance(x, (int, float)) else '')

        # Format 'Fund Name' for display
        processed_df['Reksa Dana'] = processed_df['Reksa Dana'].apply(
            lambda x: '<br>'.join(f"{i+1}. {name.strip()}" for i, name in enumerate(x.split(',')))
        )

        # Convert DataFrame to HTML without default header generation
        table_html = processed_df.to_html(classes='table table-striped table-bordered', header=False, index=False, escape=False)

        # Define custom header with rowspan and colspan
        custom_header = """
        <thead>
            <tr>
                <th rowspan="2">No.</th>
                <th rowspan="2">SID</th>
                <th rowspan="2">Nama</th>
                <th rowspan="2">Reksa Dana</th>
                <th rowspan="2">Jenis Nasabah</th>
                <th rowspan="2">Nominal</th>
                <th colspan="2">Tingkat Risiko Nasabah</th>
            </tr>
            <tr>
                <th>CDD</th>
                <th>EDD</th>
            </tr>
        </thead>
        """

        # Insert custom header into the table HTML
        table_html = table_html.replace("<tbody>", custom_header + "<tbody>")

        return render_template('result_cdd.html', tables=[table_html])
    else:
        return redirect(url_for('home'))


@app.route('/risiko_tppu', methods=['POST'])
def risiko_tppu():
    if 'username' not in session and 'admin_username' not in session:
        flash("Silakan login terlebih dahulu untuk mengakses halaman ini", "warning")
        return redirect(url_for('landingpage'))

    username = session.get('username') or session.get('admin_username', 'Unknown User')
    role = 'user' if 'username' in session else 'admin'
    log_activity(username, role, "Melakukan Proses Risiko TPPU")
    global processed_df, finalized_df


    if request.method == 'POST':
        # Upload Excel files
        file1 = request.files['file1']
        file2 = request.files['file2']
        file3 = request.files['file3']

        # Read Excel files
        df1 = pd.read_excel(file1, usecols=['SID', 'Investor Fund Unit A/C Name', 'Fund Name', 'CB Name', 'Amount (IDR Equivalent)'])
        df2 = pd.read_excel(file2, usecols=['SID', 'Investor Type', 'First Name', 'Middle Name', 'Last Name', 'ID No.', 'Occupation', 'Income Level (IDR)', 'Asset Owner', 'KTP Address', 'Correspondence Address', 'Country of Domicile', 'Country of Nationality', 'Date of Birth'])
        df3 = pd.read_excel(file3, usecols=['SID', 'Investor Type', 'Company Name', 'Income Level (IDR)', 'Country of Domicile', 'Date of Establishment'])

        # Drop rows without SID
        df1.dropna(subset=['SID'], inplace=True)
        df2.dropna(subset=['SID'], inplace=True)
        df3.dropna(subset=['SID'], inplace=True)

        # Combine first, middle, and last name into a single column
        df2['Name Nasabah'] = df2[['First Name', 'Middle Name', 'Last Name']].fillna('').agg(' '.join, axis=1)

        # Calculate age from 'Date of Birth' for individual investors
        def calculate_age(dob, investor_type):
            if pd.notnull(dob) and investor_type != 'Institusional':
                try:
                    # Parse the date format YYYYMMDD
                    dob = datetime.strptime(str(int(dob)), '%Y%m%d')
                    today = datetime.today()
                    age = today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))
                    return age
                except ValueError:
                    return None
            return None

        df2['Age'] = df2.apply(lambda row: calculate_age(row['Date of Birth'], row['Investor Type']), axis=1)

        # Calculate age from 'Date of Establishment' for institutional investors
        def calculate_institution_age(date_of_establishment):
            if pd.notnull(date_of_establishment):
                try:
                    # Parse the date format YYYYMMDD
                    establishment_date = datetime.strptime(str(int(date_of_establishment)), '%Y%m%d')
                    today = datetime.today()
                    age = today.year - establishment_date.year - ((today.month, today.day) < (establishment_date.month, establishment_date.day))
                    return age
                except ValueError:
                    return None
            return None

        # Add age for institutional investors
        df3['Age'] = df3['Date of Establishment'].apply(calculate_institution_age)

        # Create a 'Nama' column for individual or institutional investors
        df3['Nama'] = df3['Company Name']
        df2['Nama'] = df2['Name Nasabah']

        # Filter matching SID values
        df1_matched_df2 = df1[df1['SID'].isin(df2['SID'])]
        df1_matched_df3 = df1[~df1['SID'].isin(df2['SID']) & df1['SID'].isin(df3['SID'])]

        # Merge DataFrames by SID
        result_df2 = pd.merge(df1_matched_df2, df2, on='SID', how='left')
        result_df3 = pd.merge(df1_matched_df3, df3[['SID', 'Investor Type', 'Nama', 'Country of Domicile', 'Age', 'Income Level (IDR)']], on='SID', how='left')

        # Combine the results and aggregate
        final_df = pd.concat([result_df2, result_df3], ignore_index=True)
        final_df = final_df.groupby(['SID', 'Nama', 'Investor Type'], as_index=False).agg({
            'Fund Name': lambda x: ', '.join(str(i) for i in x if isinstance(i, str)),
            'Amount (IDR Equivalent)': 'sum',
            'Country of Domicile': lambda x: ', '.join(str(i) for i in x.unique() if isinstance(i, str)),
            'Age': 'first',
            'Income Level (IDR)': 'first'  # Pastikan income level diambil dari df2
        })

        # Daftar negara berisiko tinggi
        risk_countries = [
            'AF', 'AO', 'BA', 'BG', 'BF', 'CM', 'HR', 'CD', 'EC', 'GY', 'HT', 'IQ',
            'KE', 'LA', 'ML', 'MC', 'MZ', 'NA', 'NG', 'PA', 'PG', 'PH', 'SN', 'ZA',
            'SS', 'SD', 'SY', 'TZ', 'UG', 'VE', 'VN', 'YE'
        ]

        # Daftar dana berisiko tinggi
        high_risk_funds = [
            'PNM SAHAM AGRESIF', 'REKSA DANA PNM EKUITAS SYARIAH',
            'REKSA DANA PNM SAHAM UNGGULAN', 'REKSA DANA INDEKS PNM INDEKS INFOBANK15'
        ]

        # Fungsi untuk klasifikasi risiko dengan keterangan
        def classify_risk(row):
            # Klasifikasi berdasarkan negara berisiko tinggi
            if row['Country of Domicile'] in risk_countries or (
                    row['Investor Type'] == 'Individual' and row.get('Country of Nationality', '') in risk_countries
                    ):
                return 'Tinggi', 'Negara berisiko tinggi'

        # Klasifikasi berdasarkan Income Level dan Amount
            if row['Investor Type'] == 'Individual':
        # Check income level for individual investors when amount > 100 million
                if row['Amount (IDR Equivalent)'] > 100_000_000:
                    if not (row['Income Level (IDR)'] in ['> 100 - 500 million/year', '> 500 million/year']):
                        return 'Tinggi', 'Income level < Amount (IDR Equivalent)'
                elif row['Investor Type'] == 'Institusional':
        # Check income level for institutional investors when amount > 1 billion
                    if row['Amount (IDR Equivalent)'] > 500_000_000:
                        if row['Income Level (IDR)'] == '< 1 billion/year':
                            return 'Tinggi', 'Income level < Amount (IDR Equivalent)'

        # Klasifikasi berdasarkan Amount dan Usia (untuk nasabah individual)
            if row['Investor Type'] != 'Institusional':
                if pd.notnull(row['Age']) and row['Age'] > 60 and any(fund in row['Fund Name'] for fund in high_risk_funds) and row['Amount (IDR Equivalent)'] > 50_000_000:
                    return 'Tinggi', 'Usia > 60th berinvestasi (>50jt) pada reksadana high risk'
                elif row['Amount (IDR Equivalent)'] <= 100_000_000:
                    return 'Rendah', 'Amount (IDR Equivalent) < 100 jt'
                elif 100_000_000 < row['Amount (IDR Equivalent)'] <= 500_000_000:
                    return 'Sedang', 'Amount (IDR Equivalent) 100 - 500 jt'
                else:
                    return 'Tinggi', 'Amount (IDR Equivalent) > 500 jt'
            else:
        # Klasifikasi berdasarkan Amount untuk nasabah institusional
                if row['Amount (IDR Equivalent)'] <= 100_000_000:
                    return 'Rendah', 'Amount (IDR Equivalent) <100 jt'
                elif 100_000_000 < row['Amount (IDR Equivalent)'] <= 500_000_000:
                    return 'Sedang', 'Amount (IDR Equivalent) 100 - 500 jt'
                else:
                    return 'Tinggi', 'Amount (IDR Equivalent) > 500 jt'

        # Terapkan klasifikasi dan keterangan pada DataFrame
        final_df['Risiko Klasifikasi'], final_df['Keterangan'] = zip(*final_df.apply(classify_risk, axis=1))

        # Tambahkan kolom sub-risiko Rendah, Sedang, dan Tinggi
        final_df['Risiko Rendah'] = final_df['Risiko Klasifikasi'].apply(lambda x: 'X' if x == 'Rendah' else '')
        final_df['Risiko Sedang'] = final_df['Risiko Klasifikasi'].apply(lambda x: 'X' if x == 'Sedang' else '')
        final_df['Risiko Tinggi'] = final_df['Risiko Klasifikasi'].apply(lambda x: 'X' if x == 'Tinggi' else '')

        final_df['No.'] = range(1, len(final_df) + 1)
        # Update MultiIndex untuk kolom
        final_df = final_df[[ 'No.', 'SID', 'Nama', 'Investor Type', 'Fund Name', 'Amount (IDR Equivalent)',
                      'Income Level (IDR)', 'Country of Domicile', 'Age',
                      'Risiko Rendah', 'Risiko Sedang', 'Risiko Tinggi', 'Keterangan']]

        final_df.columns = pd.MultiIndex.from_tuples([
            ('No.', ''), ('SID', ''), ('Nama', ''), ('Investor Type', ''),
            ('Fund Name', ''), ('Amount (IDR Equivalent)', ''),
            ('Income Level (IDR)', ''), ('Country of Domicile', ''), ('Age', ''),
            ('Risiko Nasabah', 'Rendah'), ('Risiko Nasabah', 'Sedang'), ('Risiko Nasabah', 'Tinggi'), ('Keterangan', '')
            ])

        # Urutkan data berdasarkan risiko nasabah
        final_df.sort_values(
            by=[('Risiko Nasabah', 'Tinggi'), ('Risiko Nasabah', 'Sedang'), ('Risiko Nasabah', 'Rendah'), ('Investor Type', '')],
            ascending=[False, False, True, False],
            inplace=True
            )

        # memperbarui nomor urut lagi setelah pengurutan
        final_df['No.'] = range(1, len(final_df) + 1)

        # Tambahkan baris ringkasan jika diperlukan
        summary_row = pd.DataFrame({
            ('No.', ''): [''],
            ('SID', ''): ['TOTAL'],
            ('Nama', ''): [''],
            ('Investor Type', ''): [''],
            ('Fund Name', ''): [''],
            ('Amount (IDR Equivalent)', ''): [final_df.shape[0]],
            ('Income Level (IDR)', ''): [''],
            ('Country of Domicile', ''): [''],
            ('Age', ''): [''],
            ('Risiko Nasabah', 'Rendah'): [final_df.xs('Rendah', level=1, axis=1).value_counts().get('X', 0)],
            ('Risiko Nasabah', 'Sedang'): [final_df.xs('Sedang', level=1, axis=1).value_counts().get('X', 0)],
            ('Risiko Nasabah', 'Tinggi'): [final_df.xs('Tinggi', level=1, axis=1).value_counts().get('X', 0)],
            ('Keterangan', ''): ['']
        })

        final_df = pd.concat([final_df, summary_row], ignore_index=True)

        # Store the processed DataFrame for further use
        processed_df = final_df
        finalized_df = processed_df.copy()

        return redirect(url_for('lihat_risiko_tppu'))

    return "Gagal memproses Risiko TPPU"


@app.route('/lihat_risiko_tppu')
def lihat_risiko_tppu():
    global processed_df

    if processed_df is not None and not processed_df.empty:
        try:
            # Format 'Amount (IDR Equivalent)' for better readability
            processed_df[('Amount (IDR Equivalent)', '')] = processed_df[('Amount (IDR Equivalent)', '')].apply(
                lambda x: f'{x:,.0f}'.replace(',', '.') if isinstance(x, (int, float)) else ''
            )

            # Format 'Fund Name' for display
            processed_df[('Fund Name', '')] = processed_df[('Fund Name', '')].apply(
                lambda x: '<br>'.join(f"{i+1}. {name.strip()}" for i, name in enumerate(x.split(',')))
            )

            # Convert DataFrame to HTML but remove the default header generation
            table_html = processed_df.to_html(classes='table table-striped table-bordered', header=False, index=False, escape=False)

            # Define the custom header with rowspan and colspan
            custom_header = """
            <thead>
                <tr>
                    <th rowspan="2">No.</th>
                    <th rowspan="2">SID</th>
                    <th rowspan="2">Nama</th>
                    <th rowspan="2">Investor Type</th>
                    <th rowspan="2">Fund Name</th>
                    <th rowspan="2">Amount (IDR Equivalent)</th>
                    <th rowspan="2">Income Level (IDR)</th>
                    <th rowspan="2">Country of Domicile</th>
                    <th rowspan="2">Age</th>
                    <th colspan="3">Risiko Nasabah</th>
                    <th rowspan="2">Keterangan</th>
                </tr>
                <tr>
                    <th>Rendah</th>
                    <th>Sedang</th>
                    <th>Tinggi</th>
                </tr>

            </thead>
            """

            # Insert custom header into the table HTML
            table_html = table_html.replace("<tbody>", custom_header + "<tbody>")

            return render_template('result_risiko.html', tables=[table_html])
        except Exception as e:
            # Log error or handle it appropriately
            print(f"Error occurred: {e}")
            return render_template('home_analisis.html', error=str(e))
    else:
        return redirect(url_for('home'))

@app.route('/download_risiko_tppu')
def download_risiko_tppu():
    if 'username' not in session and 'admin_username' not in session:
        flash("Silakan login terlebih dahulu untuk mengakses halaman ini", "warning")
        return redirect(url_for('landingpage'))

    username = session.get('username') or session.get('admin_username', 'Unknown User')
    role = 'user' if 'username' in session else 'admin'
    log_activity(username, role, "Mengunduh Excel Risiko TPPU ")
    global finalized_df  # Use the locked DataFrame for download

    if finalized_df is not None:
        workbook = Workbook()
        worksheet = workbook.active

        # Format 'Fund Name' for new lines
        finalized_df['Fund Name'] = finalized_df['Fund Name'].apply(
            lambda x: '\n'.join(f"{i+1}. {name.strip()}" for i, name in enumerate(x.split(',')))
        )

        # Write DataFrame to the worksheet
        for r_idx, row in enumerate(dataframe_to_rows(finalized_df, index=False, header=True), 1):
            worksheet.append(row)

        # Adjust column widths based on the maximum length of data in each column
        columns = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I','J','K','L','M']  # Update with 'Country of Domicile'
        for col_letter in columns:
            max_length = 0
            for cell in worksheet[col_letter]:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))

            adjusted_width = max_length + 2  # Adding some padding to the width
            worksheet.column_dimensions[col_letter].width = adjusted_width

            # Special case for 'Fund Name' column (column E), setting a maximum width
            if col_letter == 'E':
                worksheet.column_dimensions[col_letter].width = 80  # Limit the width of Fund Name column

        # Apply text wrapping for 'Fund Name' column (Column E)
        for row in worksheet.iter_rows(min_row=2, min_col=5, max_col=5):
            for cell in row:
                cell.alignment = cell.alignment.copy(wrap_text=True)

        # Apply middle alignment to all columns except 'Fund Name' (column E)
        for row in worksheet.iter_rows(min_row=2):
            for cell in row:
                if cell.column_letter != 'E':  # Skip Fund Name column for middle alignment
                    cell.alignment = Alignment(horizontal='center', vertical='center')  # Middle alignment

        # Save file to BytesIO
        output = BytesIO()
        workbook.save(output)
        output.seek(0)

        return send_file(output, as_attachment=True, download_name='Result_risiko_tppu.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    else:
        return "No file to download"



@app.route('/rbs', methods=['POST'])
def rbs():
    if 'username' not in session and 'admin_username' not in session:
        flash("Silakan login terlebih dahulu untuk mengakses halaman ini", "warning")
        return redirect(url_for('landingpage'))

    username = session.get('username') or session.get('admin_username', 'Unknown User')
    role = 'user' if 'username' in session else 'admin'
    log_activity(username, role, "Melakukan Proses RBS")
    if request.method == 'POST':
        # Handle form submission for RBS
        file1 = request.files['file1']
        file2 = request.files['file2']
        file3 = request.files['file3']
        # Process the files and data here for RBS
        return "RBS processed successfully!"

# Route utama untuk memproses file SIPESAT
@app.route('/sipesat', methods=['POST'])
def sipesat():
    if 'username' not in session and 'admin_username' not in session:
        flash("Silakan login terlebih dahulu untuk mengakses halaman ini", "warning")
        return redirect(url_for('landingpage'))

    username = session.get('username') or session.get('admin_username', 'Unknown User')
    role = 'user' if 'username' in session else 'admin'
    log_activity(username, role, "Melakukan Proses SIPESAT")
    global finalized_individu, finalized_institusi

    # Fungsi untuk mendapatkan rentang tanggal berdasarkan kuartal dan tahun
    def get_date_range(quarter, year):
        if quarter == 'Q1':
            return f"{year}-01-01", f"{year}-03-31"
        elif quarter == 'Q2':
            return f"{year}-04-01", f"{year}-06-30"
        elif quarter == 'Q3':
            return f"{year}-07-01", f"{year}-09-30"
        elif quarter == 'Q4':
            return f"{year}-10-01", f"{year}-12-31"
        else:
            return None, None

    # Fungsi untuk memfilter data per kuartal dan mengambil data terbaru
    def filter_latest_data(df, date_column, start_date, end_date, date_format):
        df[date_column] = pd.to_datetime(df[date_column], format=date_format, errors='coerce')
        df = df.dropna(subset=[date_column])
        df_filtered = df[(df[date_column] >= start_date) & (df[date_column] <= end_date)]
        df_filtered = df_filtered.sort_values(by=date_column, ascending=True).drop_duplicates()
        return df_filtered

    if request.method == 'POST':
        year = request.form.get('year')
        quarter = request.form.get('quarter')

        if not year or not quarter:
            return "Tahun atau kuartal tidak valid. Silakan coba lagi."

        start_date, end_date = get_date_range(quarter, year)

        if start_date is None or end_date is None:
            return "Kuartal tidak valid. Silakan pilih kuartal yang benar."

        if 'file1' not in request.files:
            return "Tidak ada file yang diunggah. Silakan unggah file Excel dengan format .xls atau .xlsx."

        file = request.files['file1']
        if file.filename == '':
            return "Nama file kosong. Silakan pilih file untuk diunggah."

        if file:
            excel_data = pd.ExcelFile(file)
            sheet2 = excel_data.parse('INDIVIDU')
            sheet3 = excel_data.parse('INSTITUSI')

            # Filter data individu dan institusi berdasarkan kuartal dan ambil data terbaru
            finalized_individu = filter_latest_data(sheet2, 'CreateTime', start_date, end_date, "%m %d %Y %H:%M:%S")
            finalized_institusi = filter_latest_data(sheet3, 'CreateTime', start_date, end_date, "%d %m %Y %H:%M:%S")

            # Tambahkan kolom KodeNasabah dengan nilai sesuai jenis data
            finalized_individu['KodeNasabah'] = 'INDIVIDU'
            finalized_institusi['KodeNasabah'] = 'INSTITUSI'

            if 'NoPassport' not in finalized_institusi.columns:
                finalized_institusi['NoPassport'] = np.nan

            # Isi 'NoPassport' kosong dengan nilai '' saja, kolom lain dengan 'dalam proses melengkapi data'
            finalized_individu['NoPassport'] = finalized_individu['NoPassport'].fillna('')
            finalized_institusi['NoPassport'] = finalized_institusi['NoPassport'].fillna('')

            # Kolom lainnya tetap diisi 'dalam proses melengkapi data'
            finalized_individu.fillna('dalam proses melengkapi data', inplace=True)
            finalized_institusi.fillna('dalam proses melengkapi data', inplace=True)

        return redirect(url_for('lihat_sipesat'))

    return "Gagal memproses Risiko SIPESAT"


# Rute untuk menampilkan hasil di halaman web
@app.route('/lihat_sipesat')
def lihat_sipesat():
    global finalized_individu, finalized_institusi

    # Tidak mengubah 'NoPassport' yang sudah kosong, tetap kosong jika tidak ada data
    if (finalized_individu is not None and not finalized_individu.empty) or (finalized_institusi is not None and not finalized_institusi.empty):
        try:
            individu_html = finalized_individu.to_html(classes='table table-striped table-bordered text-center', index=False, escape=False) if finalized_individu is not None else ""
            institusi_html = finalized_institusi.to_html(classes='table table-striped table-bordered text-center', index=False, escape=False) if finalized_institusi is not None else ""

            return render_template('result_sipesat.html', individu_table=individu_html, institusi_table=institusi_html)
        except Exception as e:
            print(f"Error occurred: {e}")
            return "Terjadi kesalahan saat menampilkan data."
    else:
        return "Belum ada data yang diproses. Silakan unggah file di /sipesat terlebih dahulu."

@app.route('/download_sipesat')
def download_sipesat():
    if 'username' not in session and 'admin_username' not in session:
        flash("Silakan login terlebih dahulu untuk mengakses halaman ini", "warning")
        return redirect(url_for('landingpage'))

    username = session.get('username') or session.get('admin_username', 'Unknown User')
    role = 'user' if 'username' in session else 'admin'
    log_activity(username, role, "Mengunduh Excel SIPESAT")
    global finalized_individu, finalized_institusi

    # Memastikan data ada untuk diunduh
    if finalized_individu is None and finalized_institusi is None:
        return "Belum ada data yang diproses. Silakan unggah file di /sipesat terlebih dahulu."

    # Menambahkan kolom 'NoPassport' jika tidak ada di finalized_institusi
    if 'NoPassport' not in finalized_institusi.columns:
        finalized_institusi['NoPassport'] = np.nan

    # Menambahkan kolom IDPJK dengan nilai tetap 4321
    finalized_individu['IDPJK'] = 43321
    finalized_institusi['IDPJK'] = 43321

    # Mengubah ClientName menjadi huruf kapital semua
    finalized_individu['ClientName'] = finalized_individu['ClientName'].str.upper()
    finalized_institusi['ClientName'] = finalized_institusi['ClientName'].str.upper()

    # Mengubah KodeNasabah: 1 untuk Individu, 2 untuk Institusi
    finalized_individu['KodeNasabah'] = 1
    finalized_institusi['KodeNasabah'] = 2

    # Format tanggal menjadi 'Hari-Bulan-Tahun' dengan bulan singkatan Inggris
    finalized_individu['TanggalLahir'] = finalized_individu['TanggalLahir'].dt.strftime('%d-%b-%Y')
    finalized_institusi['TanggalPendirian'] = finalized_institusi['TanggalPendirian'].dt.strftime('%d-%b-%Y')

    # Memilih kolom tertentu untuk individu dan institusi
    individu_selected = finalized_individu[['IDPJK', 'KodeNasabah', 'ClientName', 'TempatLahir', 'TanggalLahir',
                                            'AlamatSuratMenyurat', 'NoKTP', 'NoPassport', 'CIFCode', 'NoNPWP']]
    institusi_selected = finalized_institusi[['IDPJK', 'KodeNasabah', 'ClientName', 'TempatPendirian', 'TanggalPendirian',
                                              'AlamatInstitusi', 'NoSKD', 'NoPassport', 'CIFCode', 'NoNPWP']]
    institusi_selected.columns = individu_selected.columns  # Menyamakan kolom institusi dengan individu

    # Menggabungkan data individu dan institusi
    combined_data = pd.concat([individu_selected, institusi_selected], ignore_index=True)
    combined_data.rename(columns={'NoPassport': 'Kolom Identitas Lain'}, inplace=True)
    combined_data['Kolom Identitas Lain'] = combined_data['Kolom Identitas Lain'].fillna('')

    # Membuat file Excel
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        workbook = writer.book
        header_format = workbook.add_format({'align': 'center', 'bold': True, 'valign': 'vcenter'})

        # Sheet 1: Gabungan data individu dan institusi dengan freeze pane di baris pertama
        combined_data.to_excel(writer, index=False, sheet_name='Gabungan')
        worksheet = writer.sheets['Gabungan']
        worksheet.freeze_panes(1, 0)  # Membekukan baris pertama
        for i, col in enumerate(combined_data.columns):
            max_width = max(combined_data[col].astype(str).map(len).max(), len(col)) + 2
            worksheet.set_column(i, i, max_width)
            worksheet.write(0, i, col, header_format)

        # Sheet 2: Data individu lengkap dengan freeze pane di baris pertama
        if finalized_individu is not None:
            finalized_individu = finalized_individu.fillna('dalam proses melengkapi data')
            finalized_individu.to_excel(writer, index=False, sheet_name='Individu Terbaru')
            worksheet = writer.sheets['Individu Terbaru']
            worksheet.freeze_panes(1, 0)  # Membekukan baris pertama
            for i, col in enumerate(finalized_individu.columns):
                max_width = max(finalized_individu[col].astype(str).map(len).max(), len(col)) + 2
                worksheet.set_column(i, i, max_width)
                worksheet.write(0, i, col, header_format)

        # Sheet 3: Data institusi lengkap dengan freeze pane di baris pertama
        if finalized_institusi is not None:
            finalized_institusi = finalized_institusi.fillna('dalam proses melengkapi data')
            finalized_institusi.to_excel(writer, index=False, sheet_name='Institusi Terbaru')
            worksheet = writer.sheets['Institusi Terbaru']
            worksheet.freeze_panes(1, 0)  # Membekukan baris pertama
            for i, col in enumerate(finalized_institusi.columns):
                max_width = max(finalized_institusi[col].astype(str).map(len).max(), len(col)) + 2
                worksheet.set_column(i, i, max_width)
                worksheet.write(0, i, col, header_format)

    # Mengirim file Excel yang dihasilkan
    output.seek(0)
    return send_file(output, as_attachment=True, download_name='sipesat_report.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')


# Rute utama untuk halaman beranda
@app.route('/Upload_Pengkinian')
def Upload_Pengkinian():
    if 'username' not in session and 'admin_username' not in session:
        flash("Silakan login terlebih dahulu untuk mengakses halaman ini", "warning")
        return redirect(url_for('landingpage'))

    username = session.get('username') or session.get('admin_username', 'Unknown User')
    role = 'user' if 'username' in session else 'admin'
    log_activity(username, role, "Melakukan Proses Upload Pengkinian")
    if 'username' not in session and 'admin_username' not in session:
        flash("Silakan login terlebih dahulu untuk mengakses halaman ini", "warning")
        return redirect(url_for('landingpage'))
    return render_template('uploadpengkinian.html')


# Temporary directory to save Excel files
TEMP_DIR = 'temp_files'
if not os.path.exists(TEMP_DIR):
    os.makedirs(TEMP_DIR)

@app.route('/Pengkinian_Data', methods=['GET', 'POST'])
def Pengkinian_Data():
    if request.method == 'POST':
        # Mengambil file yang diupload
        file_individual = request.files['file_individual']
        file_institusi = request.files['file_institusi']

        # Membaca file Excel dan menentukan tipe data untuk kolom tertentu
        if file_individual.filename.endswith('.xls'):
            df_individual = pd.read_excel(file_individual, engine='xlrd', dtype={'ID No.': str, 'SID': str, 'Date of Birth': str, 'NPWP No.': str})
        else:
            df_individual = pd.read_excel(file_individual, engine='openpyxl', dtype={'ID No.': str, 'SID': str, 'Date of Birth': str, 'NPWP No.': str})

        if file_institusi.filename.endswith('.xls'):
            df_institusi = pd.read_excel(file_institusi, engine='xlrd', dtype={'SID': str, 'Date of Establishment': str, 'NPWP No.': str})
        else:
            df_institusi = pd.read_excel(file_institusi, engine='openpyxl', dtype={'SID': str, 'Date of Establishment': str, 'NPWP No.': str})

        # Proses filtering dan renaming untuk file Individual
        df_individual_filtered = df_individual[['SID', 'ID No.', 'First Name', 'Middle Name', 'Last Name', 'Place of Birth', 'Date of Birth', 'KTP Address', 'NPWP No.']]
        df_individual_filtered['Nama Nasabah'] = df_individual_filtered[['First Name', 'Middle Name', 'Last Name']].apply(lambda x: ' '.join(x.dropna().astype(str).str.strip()), axis=1)
        df_individual_filtered = df_individual_filtered.drop(columns=['First Name', 'Middle Name', 'Last Name'])
        df_individual_filtered = df_individual_filtered.rename(columns={
            'KTP Address': 'Alamat KTP',
            'Place of Birth': 'Tempat Lahir',
            'Date of Birth': 'Tanggal Lahir',
            'ID No.': 'No. Identitas',
            'NPWP No.': 'No. NPWP'
        })

        # Pindahkan kolom 'Nama Nasabah' setelah 'SID'
        sid_index = df_individual_filtered.columns.get_loc('SID')
        df_individual_filtered.insert(sid_index + 1, 'Nama Nasabah', df_individual_filtered.pop('Nama Nasabah'))

        df_individual_filtered.fillna('Dalam proses melengkapi data', inplace=True)
        df_individual_filtered = df_individual_filtered[df_individual_filtered.isin(['Dalam proses melengkapi data']).any(axis=1)]
        df_individual_filtered = df_individual_filtered.astype(str)
        df_individual_filtered.insert(0, 'No', range(1, len(df_individual_filtered) + 1))

        # Proses filtering dan renaming untuk file Institusi
        df_institusi_filtered = df_institusi[['SID', 'Company Name', 'Place of Establishment', 'Date of Establishment', 'Company Address', 'SKD No.', 'NPWP No.']]
        df_institusi_filtered = df_institusi_filtered.rename(columns={
            'Company Name': 'Nama Perusahaan',
            'Place of Establishment': 'Tempat Pendirian',
            'Date of Establishment': 'Tanggal Pendirian',
            'Company Address': 'Alamat Institusi',
            'NPWP No.': 'No. NPWP',
            'SKD No.': 'No. SKD'
        })
        df_institusi_filtered.fillna('Dalam proses melengkapi data', inplace=True)
        df_institusi_filtered = df_institusi_filtered[df_institusi_filtered.isin(['Dalam proses melengkapi data']).any(axis=1)]
        df_institusi_filtered = df_institusi_filtered.astype(str)
        df_institusi_filtered.insert(0, 'No', range(1, len(df_institusi_filtered) + 1))

        # Simpan DataFrame ke file Excel sementara
        combined_file_path = os.path.join(TEMP_DIR, 'laporan Pengkinian Data Individual & Institusi.xlsx')

        # Menyimpan kedua DataFrame ke dalam satu file Excel dengan tabel dan border
        with pd.ExcelWriter(combined_file_path, engine='xlsxwriter') as writer:
            # Menulis DataFrame Individual ke sheet 'Data Individu'
            df_individual_filtered.to_excel(writer, sheet_name='Data Individu', index=False)
            workbook = writer.book
            worksheet = writer.sheets['Data Individu']

            # Mendapatkan range data untuk diterapkan format tabel
            individual_range = f'A1:{chr(65 + len(df_individual_filtered.columns) - 1)}{len(df_individual_filtered) + 1}'

            # Membuat format tabel dengan border
            format_border = workbook.add_format({'border': 1})
            worksheet.conditional_format(individual_range, {'type': 'no_blanks', 'format': format_border})
            worksheet.conditional_format(individual_range, {'type': 'blanks', 'format': format_border})

            # Menulis DataFrame Institusi ke sheet 'Data Institusi'
            df_institusi_filtered.to_excel(writer, sheet_name='Data Institusi', index=False)
            worksheet = writer.sheets['Data Institusi']

            # Mendapatkan range data untuk diterapkan format tabel
            institusi_range = f'A1:{chr(65 + len(df_institusi_filtered.columns) - 1)}{len(df_institusi_filtered) + 1}'

            # Membuat format tabel dengan border
            worksheet.conditional_format(institusi_range, {'type': 'no_blanks', 'format': format_border})
            worksheet.conditional_format(institusi_range, {'type': 'blanks', 'format': format_border})

        results = {
            'individual_data': df_individual_filtered.to_html(classes='data', index=False, border=0),
            'institusi_data': df_institusi_filtered.to_html(classes='data', index=False, border=0),
            'individual_file': 'laporan Pengkinian Data Individual.xlsx',
            'institusi_file': 'laporan Pengkinian Data Institusi.xlsx',
            'combined_file': 'laporan Pengkinian Data Individual & Institusi.xlsx'
        }

        return render_template('hasilpengkinian.html', results=results)

    return render_template('uploadpengkinian.html')

@app.route('/download/<filename>')
def download(filename):
    if 'username' not in session and 'admin_username' not in session:
        flash("Silakan login terlebih dahulu untuk mengakses halaman ini", "warning")
        return redirect(url_for('landingpage'))

    username = session.get('username') or session.get('admin_username', 'Unknown User')
    role = 'user' if 'username' in session else 'admin'
    log_activity(username, role, "Mengunduh Excel Pengkinian Data")
    file_path = os.path.join(TEMP_DIR, filename)
    return send_file(file_path, as_attachment=True)
@app.route('/upload_skrining', methods=['GET', 'POST'])
def upload_file_skrining():
    if 'username' not in session and 'admin_username' not in session:
        flash("Silakan login terlebih dahulu untuk mengakses halaman ini", "warning")
        return redirect(url_for('landingpage'))

    if request.method == 'POST':
        session.pop('matched_data', None)
        username = session.get('username') or session.get('admin_username', 'Unknown User')
        role = 'user' if 'username' in session else 'admin'
        log_activity(username, role, " Melakukan Proses Screening Nasabah Baru PNM")
    if request.method == 'POST':
        files = {
            'dttot': request.files.get('file_dttot'),
            'dppspm': request.files.get('file_dppspm'),
            'judionline': request.files.get('file_judionline')
        }

        for key, file in files.items():
            if file and allowed_file(file.filename):
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
                file.save(file_path)

                # Ambil file_key berdasarkan kunci
                if save_to_database(file_path, key):
                    flash(f'File {key} berhasil di-upload dan disimpan di database', 'success')
            else:
                flash(f'Format file {key} tidak diizinkan atau tidak ada file yang dipilih', 'danger')

        return redirect(url_for('view_data'))  # Redirect ke halaman data setelah sukses

    return render_template('upload_skrining.html')


def determine_file_key(filename):
    if 'DTTOT' in filename:
        return 'dttot'
    elif 'DPPSPM' in filename:
        return 'dppspm'
    elif 'JUDI_ONLINE' in filename:
        return 'judionline'
    return None


@app.route('/data', methods=['GET'])
def view_data():
    if 'username' not in session and 'admin_username' not in session:
        flash("Silakan login terlebih dahulu untuk mengakses halaman ini", "warning")
        return redirect(url_for('landingpage'))

    print(session)  # Lihat nilai session di konsol

    if request.method == 'POST':
        session.pop('matched_data', None)
        username = session.get('username') or session.get('admin_username', 'Unknown User')
        role = 'user' if 'username' in session else 'admin'
        log_activity(username, role, " View Last Data ")
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    search_query = request.args.get('search')
    print("Search Query:", search_query)  # Debugging

    dttot_data = []
    dppspm_data = []
    judionline_data = []

    # Mengambil data dari tabel dttot
    if search_query:
        cursor.execute("SELECT `Nama`, `Deskripsi`, `Alamat` FROM dttot WHERE `Nama` LIKE %s", ('%' + search_query + '%',))
    else:
        cursor.execute("SELECT `Nama`, `Deskripsi`, `Alamat` FROM dttot")

    dttot_data = cursor.fetchall()
    print("DTTOT Data:", dttot_data)  # Debugging

    # Mengambil data dari tabel dppspm
    if search_query:
        cursor.execute("SELECT `Nama`, `Alamat`, `Informasi Lain` FROM dppspm WHERE `Nama` LIKE %s", ('%' + search_query + '%',))
    else:
        cursor.execute("SELECT `Nama`, `Alamat`, `Informasi Lain` FROM dppspm")

    dppspm_data = cursor.fetchall()
    print("DPPSPM Data:", dppspm_data)  # Debugging

    # Mengambil data dari tabel judionline
    # Mengambil data dari tabel judionline
    # Mengambil data dari tabel judionline
    if search_query:
        cursor.execute("SELECT `NAMA REKENING`, `Nomor Rekening`, `Bank`, `NIK` FROM judionline WHERE `NAMA REKENING` LIKE %s", ('%' + search_query + '%',))
    else:
        cursor.execute("SELECT `NAMA REKENING`, `Nomor Rekening`, `Bank`, `NIK` FROM judionline")

    judionline_data = cursor.fetchall()
    print("Judionline Data:", judionline_data)  # Debugging


    cursor.close()
    conn.close()

    return render_template('result_skrining_nasabah.html', dttot=dttot_data, dppspm=dppspm_data, judionline=judionline_data)

import bcrypt

def check_password(user_input, stored_password_hash):
    if bcrypt.checkpw(user_input.encode('utf-8'), stored_password_hash.encode('utf-8')):
        return True
    return False





if __name__ == "__main__":
    app.run(debug=True, use_reloader=False)
    add_admin("rafli", "zidane123")
